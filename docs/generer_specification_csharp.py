# -*- coding: utf-8 -*-
"""Genere docs/Specification-fonctionnelle-PathFinder-CSharp.docx

Specification fonctionnelle complete de l'application PathFinder, implementation
C# / ASP.NET Core MVC « classique » (dossier dotnet-mvc/) :

  * vues Razor rendues cote serveur, AUCUNE API JSON, AUCUN JavaScript ;
  * deux fonctionnalites : recherche de plus court chemin sur RestitutionGraphe,
    et galerie illustree des types de graphes (images SVG construites cote serveur).

Le code inclus est repris tel quel des fichiers du depot. Les illustrations des
types de graphes proviennent de docs/img/ (voir docs/illustrations_graphes.py,
a lancer avant ce script).
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

REPO = Path(r"C:\Users\amami\GitHub\restitutiondonnees")
MVC = "dotnet-mvc"
IMG = REPO / "docs" / "img"
OUT = REPO / "docs" / "Specification-fonctionnelle-PathFinder-CSharp.docx"

doc = Document()

# ----------------------------------------------------------------------
# Styles
# ----------------------------------------------------------------------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)

cs = doc.styles.add_style("CodeBlock", 1)
cs.font.name = "Consolas"
cs.font.size = Pt(8)
cs.paragraph_format.space_before = Pt(4)
cs.paragraph_format.space_after = Pt(10)
cs.paragraph_format.left_indent = Pt(6)

for i in range(1, 4):
    h = doc.styles[f"Heading {i}"]
    h.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)


def _shade(paragraph, fill="F4F4F4"):
    # Ordre impose par le schema CT_PPr : pBdr avant shd.
    pPr = paragraph._p.get_or_add_pPr()

    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "9CA3AF")
    pBdr.append(left)
    pPr.append(pBdr)

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def code(text, caption=""):
    if caption:
        cp = doc.add_paragraph()
        cr = cp.add_run(caption)
        cr.bold = True
        cr.font.size = Pt(9)
        cp.paragraph_format.space_after = Pt(2)
    text = text.replace("\t", "    ").strip("\n")
    cp = doc.add_paragraph(style="CodeBlock")
    for i, line in enumerate(text.split("\n")):
        cp.add_run(("\n" if i else "") + line)
    _shade(cp)
    return cp


def whole(relpath):
    return (REPO / relpath).read_text(encoding="utf-8")


def image(slug, caption, width=3.2):
    path = IMG / f"{slug}.png"
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(f"[image manquante : {path.name} — lancer docs/illustrations_graphes.py]")
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)


def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)
def h3(t): doc.add_heading(t, level=3)


def p(t, bold=False, italic=False):
    par = doc.add_paragraph()
    r = par.add_run(t)
    r.bold = bold
    r.italic = italic
    return par


def bullet(items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def numbered(items):
    for it in items:
        doc.add_paragraph(it, style="List Number")


def kv(label, value):
    par = doc.add_paragraph()
    par.add_run(f"{label} : ").bold = True
    par.add_run(value)
    return par


def table(rows, font=9, header=True):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.cell(i, j)
            cell.text = str(val)
            for par in cell.paragraphs:
                for r in par.runs:
                    r.font.size = Pt(font)
                    if header and i == 0:
                        r.font.bold = True
    doc.add_paragraph()
    return t


def page_break():
    doc.add_page_break()


def add_toc():
    par = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r'TOC \o "1-3" \h \z \u')
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rpr.append(OxmlElement("w:i"))
    r.append(rpr)
    t = OxmlElement("w:t")
    t.text = "Table des matieres : dans Word, clic droit ici puis Mettre a jour les champs."
    r.append(t)
    fld.append(r)
    par._p.append(fld)


# ======================================================================
# PAGE DE GARDE
# ======================================================================
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Spécification fonctionnelle")
r.bold = True
r.font.size = Pt(26)

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("PathFinder — recherche de chemin sur RestitutionGraphe\n"
              "et galerie illustrée des types de graphes")
r.font.size = Pt(15)
r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Implémentation C# — ASP.NET Core MVC « classique »\n"
              "vues Razor rendues côté serveur · sans API JSON · sans JavaScript")
r.font.size = Pt(12)
r.italic = True

doc.add_paragraph()
doc.add_paragraph()
table([
    ["Champ", "Valeur"],
    ["Version", "3.4"],
    ["Date", "2026-08-28"],
    ["Périmètre", "dossier dotnet-mvc/ du dépôt restitutiondonnees"],
    ["Style d'implémentation", "ASP.NET Core MVC (Controllers + vues Razor .cshtml)"],
    ["Interfaces exposées", "pages HTML uniquement — aucune API, aucun endpoint JSON"],
    ["Rendu", "100 % côté serveur ; aucun JavaScript, aucun framework front, aucun npm"],
    ["Framework cible", ".NET 10 (net10.0)"],
    ["Base de données", "SQL Server — base RestitutionGraphe, table dbo.LINE_VIS_EDG"],
], font=10)

p("Évolutions : v1.0 décrivait dotnet-angular-mvc/ (API + Angular). v2.0 a "
  "ré-orienté la cible vers l'implémentation MVC pure (dotnet-mvc/). v2.1 a "
  "ajouté la comparaison de BFS avec les autres algorithmes. v3.0 place les "
  "familles de graphes et les définitions (graphe, DAG…) en tête de document "
  "(chapitre 2) et ajoute le chapitre 11, « Pistes d'amélioration de la base "
  "de données » (indexation, table d'arêtes normalisée, pré-calcul des "
  "composantes, graphe en mémoire). v3.1–3.2 détaillent le « scan » du § 11.4 "
  "(analogie des îles, code Union-Find). v3.3 implémente le § 11.5 "
  "(condensation SCC : Kosaraju, verdict d'existence orientée exact). v3.4 "
  "implémente le § 11.7 (graphe orienté chargé en mémoire, CSR, BFS sans "
  "SQL) — tout cela dans dotnet-new-scan/.", italic=True)

page_break()

# ======================================================================
h1("Table des matières")
add_toc()
page_break()

# ======================================================================
h1("1. Introduction")

h2("1.1 Objet du document")
p("Ce document décrit de façon complète le comportement fonctionnel attendu de "
  "l'application PathFinder dans son implémentation MVC pure (dossier "
  "dotnet-mvc/). Il précise le besoin, les acteurs, les cas d'usage, les règles "
  "de gestion, les routes web (aucune API), les algorithmes, les exigences non "
  "fonctionnelles, l'architecture, et il reproduit l'intégralité du code source.")

h2("1.2 Contexte et objectifs")
p("Le dépôt restitutiondonnees regroupe plusieurs implémentations d'un même "
  "domaine : la restitution (visualisation et interrogation) d'un graphe "
  "orienté de transformations de données. La table dbo.LINE_VIS_EDG peut "
  "contenir de l'ordre de 100 000 nœuds et plusieurs centaines de milliers "
  "d'arêtes.")
p("L'implémentation décrite ici est volontairement la plus classique possible : "
  "un projet ASP.NET Core MVC unique, où tout est rendu côté serveur par des "
  "vues Razor. Aucune API REST, aucun JavaScript, aucun outil de build front. "
  "Le navigateur ne reçoit que du HTML et une feuille de style.")
p("Elle rend deux services :", bold=True)
bullet([
    "Recherche de chemin : indiquer s'il existe un chemin orienté entre deux "
    "nœuds, en afficher un de longueur minimale (chaîne de nœuds + tableau "
    "détaillé par arête).",
    "Galerie des types de graphes : présenter, avec une image construite côté "
    "serveur pour chacun, les grandes familles de graphes (connexe, complet, "
    "compact, creux, pondéré, non pondéré, orienté, cyclique, arbre…) et "
    "l'algorithme de plus court chemin adapté à chacune.",
])
p("Pour situer d'emblée le vocabulaire, le chapitre 2 définit ce qu'est un "
  "graphe, ses familles et le DAG, avec une illustration par type. Le chapitre "
  "9 justifie ensuite le choix de BFS ; le chapitre 11 propose des pistes pour "
  "faire évoluer la base de données.")

h2("1.3 Périmètre")
p("Inclus :", bold=True)
bullet([
    "Page de recherche de chemin (route « / »).",
    "Galerie des types de graphes (route « /Graphes ») et génération des images "
    "SVG associées (« /Graphes/Image/{id} », « /Graphes/Build »).",
    "Accès en lecture seule à la table dbo.LINE_VIS_EDG.",
    "Cache applicatif des résultats de recherche.",
])
p("Hors périmètre :", bold=True)
bullet([
    "Toute API : pas d'endpoint JSON, pas de [ApiController], pas de route /api/*.",
    "Tout JavaScript côté client (le SVG est du balisage, pas du script).",
    "Création / modification / suppression de nœuds ou d'arêtes (lecture seule).",
    "Authentification des utilisateurs finaux.",
    "Alimentation de la base (script séparé).",
    "Recherche pondérée sur la vraie base (les arêtes de LINE_VIS_EDG n'ont pas "
    "de poids ; la notion de graphe pondéré n'est présentée qu'à titre "
    "d'illustration).",
    "Mise en œuvre des pistes du chapitre 11 (recommandations d'évolution, non "
    "implémentées).",
])

h2("1.4 Définitions et acronymes")
table([
    ["Terme", "Définition"],
    ["Nœud / sommet", "Identifiant textuel apparaissant en colonne Nodes ou "
     "NodesLie de dbo.LINE_VIS_EDG. Pas de table de nœuds dédiée."],
    ["Arête", "Lien orienté entre deux nœuds, dérivé d'une ligne selon sa "
              "colonne Direction."],
    ["Chemin", "Suite de nœuds n0…nk telle que chaque (ni, ni+1) est une arête. "
               "Longueur = k (nombre d'arêtes)."],
    ["Plus court chemin", "Chemin de longueur minimale. Non unique ; "
                          "l'application en renvoie un."],
    ["BFS", "Breadth-First Search — parcours en largeur, par cercles de "
            "distance croissante."],
    ["BFS bidirectionnel", "Deux BFS simultanés, l'un depuis la source, l'autre "
                           "depuis la cible, jusqu'à ce qu'ils se rejoignent."],
    ["DFS", "Depth-First Search — parcours en profondeur."],
    ["Composante connexe", "Sous-ensemble maximal de nœuds tous reliés entre "
     "eux (sens des arêtes ignoré : connexité faible)."],
    ["Composante fortement connexe (SCC)", "Sous-ensemble maximal de nœuds "
     "deux à deux reliés par un chemin ORIENTÉ dans chaque sens."],
    ["DAG", "Directed Acyclic Graph — graphe orienté sans cycle."],
    ["Fermeture transitive", "Relation « v est atteignable depuis u » pré-"
     "calculée pour tous les couples."],
    ["MVC", "Modèle-Vue-Contrôleur : Models/ (données), Controllers/ (routes), "
            "Views/ (Razor)."],
    ["Razor", "Moteur de vues de ASP.NET Core : fichiers .cshtml mêlant HTML et "
              "C#, rendus en HTML côté serveur."],
    ["SVG", "Scalable Vector Graphics — format d'image en XML. Balisage, pas de "
            "script."],
    ["ViewModel", "Objet simple rempli par le contrôleur et lu par la vue."],
], font=9)

h2("1.5 Documents et sources de référence")
bullet([
    "dotnet-mvc/README.md — présentation du projet.",
    "dotnet-mvc/Models/LineVisEdgRepository.cs — algorithme de référence (commenté).",
    "dotnet-mvc/Services/SvgGraphRenderer.cs — construction des images de graphes.",
    "Variante avec API + Angular : dotnet-angular-mvc/ (voir annexe 15.3).",
])

page_break()

# ======================================================================
h1("2. Notions de graphe et familles de graphes")
p("Ce chapitre pose le vocabulaire employé dans tout le document : ce qu'est "
  "un graphe, les distinctions orienté / non orienté et pondéré / non pondéré, "
  "la notion de DAG, puis un panorama illustré des dix grandes familles de "
  "graphes. Le choix de l'algorithme est traité plus loin (chapitre 9).")

h2("2.1 Définition d'un graphe")
p("Un graphe G est un couple (V, E) où :")
bullet([
    "V (vertices, sommets) est un ensemble fini d'éléments appelés nœuds ou "
    "sommets ;",
    "E (edges, arêtes) est un ensemble de liens entre deux sommets : une paire "
    "{u, v} si le graphe est non orienté, un couple (u, v) s'il est orienté.",
])
p("On note n = |V| l'ordre du graphe (nombre de sommets) et m = |E| sa taille "
  "(nombre d'arêtes). Un graphe modélise « des objets et des liens entre eux » : "
  "ici, des données (nœuds) et les transformations qui produisent une donnée à "
  "partir d'une autre (arêtes).")
p("Vocabulaire :", bold=True)
table([
    ["Terme", "Définition"],
    ["Voisin de u", "Sommet v relié à u par une arête."],
    ["Degré de u", "Nombre d'arêtes incidentes à u. En orienté on distingue le "
     "degré entrant (arêtes vers u) et le degré sortant (arêtes issues de u)."],
    ["Chemin", "Suite de sommets s0, s1, …, sk telle que chaque (si, si+1) est "
     "une arête. Longueur = k (nombre d'arêtes)."],
    ["Chemin simple", "Chemin sans sommet répété."],
    ["Cycle", "Chemin d'au moins une arête revenant à son sommet de départ."],
    ["Distance de u à v", "Longueur du plus court chemin de u à v (infinie s'il "
     "n'en existe aucun)."],
    ["Graphe connexe", "Graphe non orienté où il existe un chemin entre toute "
     "paire de sommets."],
    ["Composante connexe", "Sous-ensemble maximal de sommets deux à deux reliés."],
    ["Fortement connexe", "Graphe orienté où, pour toute paire (u, v), il "
     "existe un chemin orienté de u à v ET de v à u."],
    ["Graphe creux / dense", "Creux : m de l'ordre de n. Dense (« compact ») : "
     "m de l'ordre de n², proche du maximum n(n-1)/2."],
], font=9)

h2("2.2 Graphe orienté et non orienté")
p("Dans un graphe non orienté, une arête {u, v} se parcourt dans les deux "
  "sens. Dans un graphe orienté (digraphe), une arête (u, v) va de u vers v et "
  "pas l'inverse : (u, v) et (v, u) sont deux arêtes distinctes. La recherche "
  "de chemin de l'application est orientée — elle suit les arêtes dans leur "
  "sens (colonne Direction de LINE_VIS_EDG, voir § 4.3).")

h2("2.3 Graphe pondéré et non pondéré")
p("Un graphe pondéré associe à chaque arête un poids numérique (distance, "
  "coût, durée…). Le « plus court chemin » y minimise la somme des poids.")
p("Un graphe non pondéré n'a pas de poids : toutes les arêtes se valent. La "
  "longueur d'un chemin est son nombre d'arêtes ; le plus court chemin est "
  "celui qui en compte le moins. C'est le cas ici : LINE_VIS_EDG ne stocke "
  "aucun poids d'arête.")
p("Conséquence essentielle : sur un graphe non pondéré, le parcours en largeur "
  "(BFS) visite les sommets par distance croissante depuis la source. Le "
  "premier instant où il atteint la cible correspond donc nécessairement à un "
  "plus court chemin. Cette propriété tombe dès qu'il y a des poids (il faut "
  "alors Dijkstra — voir § 9.4).")

h2("2.4 Graphe orienté sans cycle (DAG)")
p("Un DAG (Directed Acyclic Graph, graphe orienté sans cycle) est un graphe "
  "orienté dans lequel il n'existe aucun cycle : en suivant les arêtes dans "
  "leur sens, on ne peut jamais revenir à un sommet déjà rencontré.")
p("Propriétés d'un DAG :", bold=True)
bullet([
    "Il admet au moins un tri topologique : un classement linéaire des sommets "
    "tel que toute arête (u, v) va d'un u placé avant v. Il y en a souvent "
    "plusieurs.",
    "Il possède au moins une source (degré entrant nul) et au moins un puits "
    "(degré sortant nul).",
    "On peut affecter à chaque sommet un niveau = longueur du plus long chemin "
    "depuis une source ; c'est ce que fait SvgGraphRenderer pour dessiner un "
    "DAG par rangées (§ 10.2).",
    "Le plus court (et le plus long) chemin dans un DAG se calcule en temps "
    "linéaire O(n + m) en parcourant les sommets dans l'ordre topologique — "
    "plus simple que Dijkstra, et valable même avec des poids négatifs.",
    "Contracter les cycles d'un graphe orienté quelconque (ses composantes "
    "fortement connexes) produit toujours un DAG : le « graphe condensé » "
    "(utilisé au § 11.5).",
])
p("Un arbre enraciné est un cas particulier de DAG. La table LINE_VIS_EDG est "
  "orientée mais n'est PAS garantie acyclique : des cycles peuvent exister "
  "(voir § 2.5.9), ce qui oblige à marquer les sommets déjà visités pendant le "
  "parcours.")

h2("2.5 Panorama illustré des familles de graphes")
p("Chaque famille est illustrée par un petit graphe d'exemple, défini en dur "
  "dans dotnet-mvc/Models/GraphSamples.cs et rendu en image par la même logique "
  "que la galerie de l'application (chapitre 10). Les illustrations de ce "
  "document sont produites à l'identique par docs/illustrations_graphes.py.",
  italic=True)

h3("2.5.1 Graphe connexe")
image("connexe", "Graphe connexe — une seule composante, tous les nœuds reliés")
p("Depuis n'importe quel nœud on atteint tous les autres. La recherche de "
  "chemin réussit toujours entre deux nœuds d'un même graphe connexe non "
  "orienté.")

h3("2.5.2 Graphe non connexe")
image("non-connexe", "Graphe non connexe — trois composantes {A,B,C}, {D,E}, {F,G}")
p("Aucune arête entre les composantes : un chemin entre deux nœuds de "
  "composantes différentes n'existe pas. C'est le cas « aucun chemin » le plus "
  "fréquent — et celui que le pré-calcul des composantes (§ 11.4) permet de "
  "trancher instantanément.")

h3("2.5.3 Graphe complet (K5)")
image("complet", "Graphe complet K5 — les 10 arêtes possibles entre 5 nœuds")
p("Toute paire de nœuds est reliée : n(n-1)/2 arêtes. Le plus court chemin "
  "entre deux nœuds fait toujours une seule arête.")

h3("2.5.4 Graphe compact (dense)")
image("compact", "Graphe compact — 12 arêtes pour 6 nœuds, proche du complet")
p("Beaucoup d'arêtes (de l'ordre de n²). Une matrice d'adjacence devient un "
  "stockage raisonnable ; le BFS explore vite un grand voisinage.")

h3("2.5.5 Graphe creux (sparse)")
image("creux", "Graphe creux — 6 arêtes pour 7 nœuds")
p("Peu d'arêtes (de l'ordre de n). C'est le profil de la base "
  "RestitutionGraphe (2 à 6 arêtes par nœud). On stocke en listes d'adjacence, "
  "et le BFS par paliers reste peu coûteux par niveau.")

h3("2.5.6 Graphe non pondéré")
image("non-pondere", "Graphe non pondéré — les arêtes n'ont pas de valeur")
p("Les arêtes indiquent seulement l'existence d'un lien ; la longueur d'un "
  "chemin est son nombre d'arêtes. C'est le modèle de cette application : un "
  "BFS (ici bidirectionnel) donne le plus court chemin. Justification complète "
  "au chapitre 9.")

h3("2.5.7 Graphe pondéré")
image("pondere", "Graphe pondéré — chaque arête porte un poids")
p("Le plus court chemin minimise la somme des poids, pas le nombre d'arêtes. "
  "Un BFS ne suffit plus : il faut Dijkstra (poids positifs) ou Bellman-Ford "
  "(voir § 9.4). Hors périmètre de l'application, montré pour situer le cas non "
  "pondéré par contraste.")

h3("2.5.8 Graphe orienté sans cycle (DAG)")
image("oriente-dag", "DAG — arêtes fléchées, aucun cycle, niveaux visibles")
p("Les arêtes ont un sens et il n'y a aucun cycle. Les nœuds s'ordonnent par "
  "niveaux (tri topologique). Définition et propriétés au § 2.4. La table "
  "LINE_VIS_EDG est orientée mais pas garantie acyclique ; le BFS respecte le "
  "sens des arêtes.")

h3("2.5.9 Graphe orienté cyclique")
image("cyclique", "Graphe cyclique — le cycle A → B → C → A")
p("En suivant les arêtes on peut revenir à son point de départ. Le parcours "
  "doit marquer les nœuds déjà visités (rôle des dictionnaires forwardPrev / "
  "backwardNext, § 8.2) pour ne pas boucler.")

h3("2.5.10 Arbre")
image("arbre", "Arbre — connexe, sans cycle, n-1 arêtes")
p("Il existe exactement un chemin entre deux nœuds quelconques. Le plus court "
  "chemin est donc l'unique chemin ; un simple parcours suffit.")

h2("2.6 Tableau de synthèse des familles")
table([
    ["Type", "Connexe", "Orienté", "Pondéré", "Cycle", "Arêtes", "Plus court chemin"],
    ["connexe", "oui", "non", "non", "oui", "~n", "BFS"],
    ["non connexe", "non", "non", "non", "oui", "~n", "BFS (par composante)"],
    ["complet Kn", "oui", "non", "non", "oui", "n(n-1)/2", "BFS (1 arête)"],
    ["compact / dense", "oui", "non", "non", "oui", "~n²", "BFS"],
    ["creux / sparse", "souvent", "non", "non", "possible", "~n", "BFS"],
    ["non pondéré", "—", "—", "non", "—", "—", "BFS  ← cas de l'appli"],
    ["pondéré", "—", "—", "oui", "—", "—", "Dijkstra / Bellman-Ford"],
    ["DAG", "possible", "oui", "non", "non", "variable", "BFS / tri topologique"],
    ["cyclique orienté", "possible", "oui", "non", "oui", "variable", "BFS avec nœuds visités"],
    ["arbre", "oui", "non", "non", "non", "n-1", "unique chemin"],
], font=8)

page_break()

# ======================================================================
h1("3. Description générale")

h2("3.1 Besoin métier")
p("Un analyste consultant le graphe des transformations doit pouvoir déterminer "
  "si une donnée (nœud source) en alimente une autre (nœud cible), directement "
  "ou indirectement, et par quel enchaînement le plus court. La volumétrie "
  "interdit de charger le graphe entier : la réponse est calculée par des "
  "requêtes ciblées. La galerie, elle, répond à un besoin pédagogique : situer "
  "le graphe manipulé (orienté, creux, non pondéré) parmi les familles "
  "classiques et comprendre pourquoi un simple BFS suffit ici.")

h2("3.2 Acteurs")
table([
    ["Acteur", "Rôle", "Interactions"],
    ["Utilisateur / analyste", "Recherche un chemin ; consulte la galerie.",
     "Pages « / » et « /Graphes »."],
    ["Exploitant", "Déploie, configure, génère les fichiers images.",
     "Variables d'environnement ; route « /Graphes/Build »."],
    ["SQL Server (RestitutionGraphe)", "Héberge les données du graphe.",
     "Requêtes en lecture émises par le Model."],
], font=9)

h2("3.3 Architecture fonctionnelle")
p("Un unique process ASP.NET Core. Le pipeline HTTP :")
numbered([
    "UseStaticFiles sert wwwroot/ (la feuille de style, et les images SVG une "
    "fois construites).",
    "La route MVC par défaut {controller=Home}/{action=Index}/{id?} dirige "
    "vers un contrôleur.",
    "HomeController.Index rend la page de recherche (et interroge le Model + le "
    "cache si une recherche est demandée).",
    "GraphesController rend la galerie et construit les images (via "
    "SvgGraphRenderer).",
])
code(
    "Navigateur\n"
    "   |  GET /                      -> page HTML : formulaire + resultat\n"
    "   |  GET /Graphes               -> page HTML : galerie de vignettes\n"
    "   |  GET /Graphes/Image/arbre   -> image  (Content-Type: image/svg+xml)\n"
    "   |  GET /Graphes/Build         -> page HTML : recap des fichiers ecrits\n"
    "   v\n"
    "+------------------------- ASP.NET Core (un seul port) -------------------------+\n"
    "|  UseStaticFiles ................ wwwroot/ (css, img/graphes/*.svg)            |\n"
    "|  MVC route par defaut                                                        |\n"
    "|     HomeController.Index ....... Views/Home/Index.cshtml                      |\n"
    "|        -> IMemoryCache -> LineVisEdgRepository.ShortestPath / DescribePath    |\n"
    "|     GraphesController.Index .... Views/Graphes/Index.cshtml                   |\n"
    "|     GraphesController.Image .... SvgGraphRenderer.Render -> FileResult (svg)  |\n"
    "|     GraphesController.Build .... ecrit wwwroot/img/graphes/*.svg              |\n"
    "+---------------------------------------------------------------------------+\n"
    "                         |  requetes SQL parametrees (lecture seule)\n"
    "                         v\n"
    "            SQL Server — base RestitutionGraphe / dbo.LINE_VIS_EDG",
    "Figure 1 — Architecture fonctionnelle (runtime)")

h2("3.4 Environnement technique")
table([
    ["Élément", "Choix"],
    ["Runtime", ".NET 10 (net10.0), ASP.NET Core"],
    ["Style", "MVC — AddControllersWithViews, vues Razor"],
    ["Accès données", "Microsoft.Data.SqlClient 7.0.2, ADO.NET synchrone"],
    ["Cache", "IMemoryCache, SizeLimit 10 000 entrées"],
    ["Rendu graphes", "SvgGraphRenderer — SVG construit à la main, sans dépendance"],
    ["Frontend", "aucun — HTML/CSS servis tels quels, pas de JavaScript"],
    ["Hébergement", "process unique, port HTTP 5175 (launchSettings.json)"],
    ["Auth SQL", "Trusted_Connection (identité Windows du process)"],
], font=9)

h2("3.5 Positionnement par rapport à dotnet-angular-mvc/")
table([
    ["Aspect", "dotnet-mvc/ (ce document)", "dotnet-angular-mvc/"],
    ["Affichage", "vues Razor, HTML/CSS", "Angular + Cytoscape (JavaScript)"],
    ["Interfaces", "pages HTML uniquement", "API JSON /api/path, /api/health"],
    ["Requête utilisateur", "rechargement de page (form GET)", "fetch AJAX"],
    ["Build front", "aucun", "npm install + ng build"],
    ["Visualisation du chemin", "chaîne de nœuds + tableau détaillé",
     "graphe interactif Cytoscape"],
    ["Cœur métier", "LineVisEdgRepository (BFS bidi) + DescribePath",
     "LineVisEdgRepository (BFS bidi)"],
], font=8.5)

page_break()

# ======================================================================
h1("4. Modèle de données")

h2("4.1 La table dbo.LINE_VIS_EDG")
p("Unique source de vérité. Chaque ligne exprime une relation orientée entre "
  "deux nœuds.")
table([
    ["Colonne", "Type", "Description"],
    ["Nodes", "VARCHAR(8000)", "Identifiant d'un nœud (extrémité 1)."],
    ["Direction", "chaîne", "'predecesseur' ou 'successeur' — rôle de Nodes "
     "vis-à-vis de NodesLie (voir 4.3)."],
    ["NodesLie", "VARCHAR(8000)", "Identifiant de l'autre nœud (extrémité 2)."],
    ["Transformation", "chaîne, nullable", "Nature de la transformation portée "
     "par l'arête (SELECT, JOIN…). Affichée dans le tableau détaillé du chemin."],
], font=9)

h2("4.2 Notion de nœud")
p("Un nœud n'a pas d'enregistrement propre : il existe dès qu'il apparaît au "
  "moins une fois en colonne Nodes ou NodesLie (règle RG-01).")

h2("4.3 Sémantique de Direction")
table([
    ["Direction", "Signification", "Arête orientée dérivée"],
    ["predecesseur", "Nodes précède NodesLie", "Nodes → NodesLie"],
    ["successeur", "Nodes suit NodesLie", "NodesLie → Nodes"],
], font=9.5)

h2("4.4 Règles de dérivation d'une arête")
table([
    ["On cherche", "Lignes concernées", "Interprétation"],
    ["arêtes SORTANTES de X", "Nodes = X ET Direction = 'predecesseur'", "X → NodesLie"],
    ["", "NodesLie = X ET Direction = 'successeur'", "X → Nodes"],
    ["arêtes ENTRANTES de X", "NodesLie = X ET Direction = 'predecesseur'", "Nodes → X"],
    ["", "Nodes = X ET Direction = 'successeur'", "NodesLie → X"],
], font=9)
p("Ces quatre cas fondent FetchEdgesFrom (front avant du BFS) et FetchEdgesInto "
  "(front arrière). Chaque cas correspond à un index composite distinct, d'où "
  "deux requêtes séparées plutôt qu'un OR. Le § 11.3 propose de matérialiser "
  "une table d'arêtes déjà orientée pour supprimer cette dérivation.")

h2("4.5 Contrainte de typage — VARCHAR(8000) explicite")
p("Nodes et NodesLie sont des colonnes VARCHAR. Par défaut Microsoft.Data."
  "SqlClient envoie une chaîne .NET en paramètre NVARCHAR ; comparer VARCHAR à "
  "NVARCHAR force SQL Server à convertir la colonne, ce qui interdit l'usage de "
  "l'index (balayage complet). Le code type donc explicitement chaque paramètre "
  "en SqlDbType.VarChar(8000) — méthode AddVarChar (RG-11).")

page_break()

# ======================================================================
h1("5. Exigences fonctionnelles — cas d'usage")

h2("5.1 UC-01 — Rechercher un chemin entre deux nœuds")
kv("Acteur", "Utilisateur / analyste")
kv("Objectif", "Savoir s'il existe un chemin orienté de la source vers la "
   "cible et en obtenir un de longueur minimale.")
kv("Préconditions", "Le service est démarré et joint RestitutionGraphe.")
kv("Déclencheur", "L'utilisateur saisit source et cible dans le formulaire de "
   "la page « / » et le soumet (bouton « Chercher » ou touche Entrée). Le "
   "formulaire est un <form method=\"get\"> : la page se recharge sur "
   "/?source=…&target=….")
p("Scénario nominal :", bold=True)
numbered([
    "Le contrôleur reçoit source, target, maxDepth (défaut 12) depuis la "
    "chaîne de requête.",
    "Il supprime les espaces de début/fin des deux valeurs.",
    "Il plafonne maxDepth à 20 (RG-08) et construit la clé de cache "
    "path:{source}:{target}:{maxDepth} (RG-09).",
    "En cas d'absence dans le cache : vérification de l'existence des deux "
    "nœuds (RG-01), puis BFS bidirectionnel par paliers (chapitre 8), borné "
    "par maxDepth et 30 000 nœuds par sens (RG-10).",
    "Le résultat (trouvé ou non) est mémorisé 5 minutes.",
    "Si un chemin est trouvé et comporte au moins une arête, le contrôleur "
    "relit pour chaque arête consécutive la Transformation associée "
    "(DescribePath).",
    "La vue rend une page HTML complète : bandeau vert « ✓ Un chemin existe "
    "(n arêtes) », chaîne des nœuds, puis tableau détaillé (# / De / Vers / "
    "Transformation).",
])
p("Scénarios alternatifs et exceptions :", bold=True)
table([
    ["Réf.", "Condition", "Comportement"],
    ["A1", "source = target (nœud existant)", "Page avec « Un chemin existe "
     "(0 arête) » et le message « Chemin de longueur 0 : la source et la cible "
     "sont le même nœud ». RG-02."],
    ["A2", "source ou target absent de la base", "Bandeau rouge « ✗ Aucun "
     "chemin de … vers … ». RG-03."],
    ["A3", "aucun chemin en maxDepth paliers", "Même bandeau rouge qu'en A2. RG-04."],
    ["A4", "un des deux champs vide", "Bandeau « Renseigne un nœud source ET "
     "un nœud cible ». Aucun accès base. RG-05."],
    ["A5", "les deux champs vides (arrivée sur « / »)", "Page avec le seul "
     "formulaire, sans message. "],
    ["A6", "base injoignable", "Page d'erreur ASP.NET Core standard (500)."],
])
kv("Postconditions", "Aucune donnée modifiée. Résultat en cache 5 min. L'URL "
   "de la page porte les paramètres — elle est partageable / rechargeable.")

h2("5.2 UC-02 — Consulter le chemin trouvé")
kv("Acteur", "Utilisateur / analyste")
kv("Préconditions", "UC-01 a renvoyé un chemin.")
p("La vue affiche, l'une sous l'autre :")
numbered([
    "La chaîne ordonnée des nœuds, séparés par « → », premier nœud en bleu, "
    "dernier en orange.",
    "Un tableau HTML, une ligne par arête : numéro, nœud de départ, nœud "
    "d'arrivée, transformation (ou « — » si la ligne n'en porte pas).",
])
p("Tout est statique : pas de zoom, pas d'interaction, pas de requête "
  "supplémentaire. Une nouvelle recherche recharge la page.")

h2("5.3 UC-03 — Réinitialiser la recherche")
kv("Déclencheur", "Clic sur le lien « Effacer » (pointe vers « / » sans "
   "paramètre).")
p("La page « / » se recharge vide : champs vides, ni bandeau ni résultat.")

h2("5.4 UC-04 — Consulter la galerie des types de graphes")
kv("Acteur", "Utilisateur / analyste")
kv("Déclencheur", "Accès à « /Graphes » (lien de navigation).")
p("Scénario nominal :", bold=True)
numbered([
    "Le contrôleur passe le catalogue GraphSamples.All à la vue.",
    "La vue affiche une vignette par type : une image (balise <img> pointant "
    "vers /Graphes/Image/{id}), le libellé, les caractéristiques (nombre de "
    "nœuds et d'arêtes, orienté ou non, pondéré ou non) et une description.",
    "Chaque image est un SVG construit à la volée par le serveur (UC-05).",
])
kv("Règle associée", "RG-15.")

h2("5.5 UC-05 — Obtenir l'image d'un type de graphe")
kv("Acteur", "Le navigateur (via les <img> de la galerie), ou l'utilisateur "
   "directement.")
kv("Déclencheur", "GET /Graphes/Image/{id} où id est le slug d'un type "
   "(ex. « arbre », « pondere »).")
p("Scénario nominal :", bold=True)
numbered([
    "Le contrôleur recherche le GraphSample de slug id.",
    "S'il n'existe pas : réponse 404.",
    "Sinon SvgGraphRenderer calcule la disposition des nœuds (cercle, ou "
    "niveaux pour un arbre / DAG) et produit le document SVG.",
    "Réponse : le SVG, Content-Type image/svg+xml. Ce n'est pas du JSON : "
    "c'est une image (FileResult).",
])

h2("5.6 UC-06 — Construire les fichiers images")
kv("Acteur", "Exploitant")
kv("Déclencheur", "GET /Graphes/Build.")
p("Scénario nominal :", bold=True)
numbered([
    "Le contrôleur crée le dossier wwwroot/img/graphes/ si besoin.",
    "Pour chaque type du catalogue, il écrit <slug>.svg (contenu = sortie de "
    "SvgGraphRenderer).",
    "La vue affiche un bandeau de confirmation et la liste des fichiers "
    "produits, chacun étant un lien vers le fichier statique.",
])
kv("Postconditions", "Les fichiers wwwroot/img/graphes/*.svg existent et sont "
   "servables directement par UseStaticFiles. RG-16.")

page_break()

# ======================================================================
h1("6. Règles de gestion")
table([
    ["Réf.", "Règle"],
    ["RG-01", "Un nœud existe s'il apparaît dans au moins une ligne, en "
     "colonne Nodes ou NodesLie."],
    ["RG-02", "Si source = target et que le nœud existe : chemin de longueur 0, "
     "path = [source]."],
    ["RG-03", "Si la source ou la cible n'existe pas : « aucun chemin », sans "
     "exécuter de parcours."],
    ["RG-04", "Si aucun chemin n'est trouvé en au plus maxDepth paliers : "
     "« aucun chemin »."],
    ["RG-05", "La recherche n'est pas lancée si l'un des deux champs est vide "
     "après suppression des espaces."],
    ["RG-06", "Le formulaire est soumis en GET : les paramètres figurent dans "
     "l'URL du résultat."],
    ["RG-07", "Le « plus court chemin » se mesure en nombre d'arêtes ; les "
     "arêtes de LINE_VIS_EDG ne portent pas de poids."],
    ["RG-08", "maxDepth est plafonné : effectiveMaxDepth = min(maxDepth, 20). "
     "Défaut 12."],
    ["RG-09", "Le résultat (trouvé ou non) est mis en cache 5 minutes sous la "
     "clé path:{source}:{target}:{effectiveMaxDepth} ; cache borné à 10 000 "
     "entrées."],
    ["RG-10", "Chaque sens du parcours cesse de s'étendre au-delà de 30 000 "
     "nœuds visités."],
    ["RG-11", "Tout paramètre comparé à Nodes/NodesLie est typé "
     "SqlDbType.VarChar(8000)."],
    ["RG-12", "Les requêtes sur une liste de nœuds sont découpées en lots de "
     "1000 paramètres au plus."],
    ["RG-13", "Le tableau détaillé du chemin affiche une ligne par arête "
     "consécutive, avec la Transformation lue dans LINE_VIS_EDG (ou « — »)."],
    ["RG-14", "Aucune réponse de l'application n'est au format JSON ; aucune "
     "route ne commence par /api. Les seules sorties sont des pages HTML et, "
     "pour /Graphes/Image, une image SVG."],
    ["RG-15", "La galerie /Graphes présente exactement les types du catalogue "
     "GraphSamples.All, dans l'ordre déclaré."],
    ["RG-16", "/Graphes/Build (re)génère wwwroot/img/graphes/<slug>.svg pour "
     "tous les types ; il écrase les fichiers existants."],
], font=9)

page_break()

# ======================================================================
h1("7. Interface web — routes MVC")
p("Il n'y a pas d'API. Les routes ci-dessous rendent des pages HTML (ou, pour "
  "une seule d'entre elles, une image). Toutes sont en méthode GET, sans "
  "authentification.")

h2("7.1 GET / — page de recherche de chemin")
p("Contrôleur : HomeController.Index.")
table([
    ["Paramètre (query)", "Type", "Oblig.", "Défaut", "Rôle"],
    ["source", "string", "non*", "—", "nœud de départ"],
    ["target", "string", "non*", "—", "nœud d'arrivée"],
    ["maxDepth", "int", "non", "12", "paliers max (plafonné à 20)"],
], font=8.5)
p("* Sans source ni target : page vierge (formulaire seul). Avec un seul des "
  "deux : message d'erreur de saisie.")
p("États de la page rendus (voir UC-01) :", bold=True)
bullet([
    "vierge — formulaire seul ;",
    "erreur de saisie — bandeau « Renseigne un nœud source ET un nœud cible » ;",
    "chemin trouvé — bandeau vert + chaîne de nœuds + tableau détaillé ;",
    "chemin de longueur 0 — bandeau vert + note ;",
    "aucun chemin — bandeau rouge.",
])

h2("7.2 GET /Graphes — galerie des types de graphes")
p("Contrôleur : GraphesController.Index. Aucun paramètre. Rend une grille de "
  "vignettes (une par type), chacune avec son image, ses caractéristiques et sa "
  "description.")

h2("7.3 GET /Graphes/Image/{id} — image d'un type de graphe")
p("Contrôleur : GraphesController.Image. Le segment de route {id} est le slug "
  "du type (connexe, non-connexe, complet, compact, creux, non-pondere, "
  "pondere, oriente-dag, cyclique, arbre).")
table([
    ["Cas", "Réponse"],
    ["slug connu", "200, Content-Type image/svg+xml, corps = document SVG"],
    ["slug inconnu", "404"],
], font=9.5)

h2("7.4 GET /Graphes/Build — génération des fichiers SVG")
p("Contrôleur : GraphesController.Build. Écrit wwwroot/img/graphes/<slug>.svg "
  "pour tous les types, puis rend une page listant les fichiers écrits. "
  "Idempotent (réécrit à chaque appel).")

page_break()

# ======================================================================
h1("8. Algorithme — BFS bidirectionnel")

h2("8.1 Principe")
p("La recherche exécute deux parcours en largeur simultanés : un front avant "
  "qui part de la source et suit les arêtes dans leur sens (arêtes sortantes), "
  "un front arrière qui part de la cible et remonte les arêtes (arêtes "
  "entrantes). Les fronts progressent d'un palier chacun à leur tour ; dès "
  "qu'un nœud est atteint des deux côtés, un plus court chemin est trouvé.")

h2("8.2 Structures de données")
table([
    ["Nom", "Type", "Rôle"],
    ["forwardPrev", "Dictionnaire nœud → nœud", "prédécesseur de chaque nœud "
     "atteint depuis la source (source : null)."],
    ["backwardNext", "Dictionnaire nœud → nœud", "successeur de chaque nœud "
     "atteint depuis la cible (cible : null)."],
    ["forwardFrontier / backwardFrontier", "Listes de nœuds", "frontières "
     "courantes des deux fronts."],
    ["maxDepth", "entier ≤ 20", "nombre maximal de paliers."],
    ["maxVisitedPerSide", "30 000", "plafond de nœuds par sens."],
], font=9)

h2("8.3 Déroulé pas à pas")
numbered([
    "Vérifier l'existence de source et cible (sinon : non trouvé).",
    "Si source = cible : renvoyer [source].",
    "Initialiser les dictionnaires et frontières.",
    "Répéter au plus maxDepth fois : étendre le front avant aux paliers pairs, "
    "le front arrière aux paliers impairs.",
    "Front avant : récupérer les arêtes sortantes de la frontière "
    "(FetchEdgesFrom) ; pour chaque (u → v) avec u connu et v inconnu : "
    "forwardPrev[v] = u ; si v ∈ backwardNext → rencontre.",
    "Front arrière : symétrique avec FetchEdgesInto ; pour chaque (u → v) avec "
    "v connu et u inconnu : backwardNext[u] = v ; si u ∈ forwardPrev → "
    "rencontre.",
    "À la rencontre : recoller « source → point de rencontre » (via "
    "forwardPrev) et « point de rencontre → cible » (via backwardNext).",
    "Après maxDepth paliers, ou frontières vides : non trouvé.",
])

h2("8.4 Pseudo-code")
code(
    "fonction PlusCourtChemin(source, cible, maxDepth <= 20):\n"
    "    si non Existe(source) ou non Existe(cible): retourner NON_TROUVE\n"
    "    si source == cible: retourner [source]\n"
    "    forwardPrev  <- { source: NUL } ; forwardFrontier  <- [source]\n"
    "    backwardNext <- { cible:  NUL } ; backwardFrontier <- [cible]\n"
    "    pour step de 0 a maxDepth-1:\n"
    "        si forwardFrontier vide et backwardFrontier vide: arreter\n"
    "        si (step pair) et forwardFrontier non vide et |forwardPrev| < 30000:\n"
    "            pour chaque (u -> v) dans FetchEdgesFrom(forwardFrontier):\n"
    "                si u pas dans forwardPrev ou v deja dans forwardPrev: continuer\n"
    "                forwardPrev[v] <- u\n"
    "                si v dans backwardNext: retourner Recoller(v)\n"
    "            forwardFrontier <- nouveaux v\n"
    "        sinon si backwardFrontier non vide et |backwardNext| < 30000:\n"
    "            pour chaque (u -> v) dans FetchEdgesInto(backwardFrontier):\n"
    "                si v pas dans backwardNext ou u deja dans backwardNext: continuer\n"
    "                backwardNext[u] <- v\n"
    "                si u dans forwardPrev: retourner Recoller(u)\n"
    "            backwardFrontier <- nouveaux u\n"
    "    retourner NON_TROUVE",
    "Pseudo-code de référence")

h2("8.5 Exemple déroulé")
p("Graphe : N1→N2, N2→N3, N3→N4, N1→N9, N9→N3. Recherche N1 → N4.")
table([
    ["Palier", "Front", "Frontière", "Découvertes", "Rencontre ?"],
    ["0", "avant", "[N1]", "N2 (prev N1), N9 (prev N1)", "non"],
    ["1", "arrière", "[N4]", "N3 (next N4)", "non"],
    ["2", "avant", "[N2, N9]", "N3 (prev N2) — déjà dans backwardNext", "OUI sur N3"],
], font=8.5)
p("Chemin reconstitué : [N1, N2, N3, N4], longueur 3.")

h2("8.6 Complexité")
bullet([
    "BFS à sens unique : ≈ d^R nœuds visités (d = degré moyen, R = longueur du "
    "chemin).",
    "BFS bidirectionnel : ≈ 2·d^(R/2) — gain exponentiel en R.",
    "Coût dominé par les allers-retours SQL : au plus 2·maxDepth lots de "
    "requêtes, chacun découpé en sous-requêtes de 1000 paramètres. Le "
    "chapitre 11 explique comment supprimer cette latence.",
])

h2("8.7 Enrichissement du chemin — DescribePath")
p("Spécifique à cette implémentation MVC : une fois le chemin obtenu, pour "
  "chaque arête consécutive (u, v) le Model relit la Transformation de la ligne "
  "correspondante — essai des deux formes de stockage : (Nodes=u, predecesseur, "
  "NodesLie=v) puis (Nodes=v, successeur, NodesLie=u). Le résultat alimente le "
  "tableau détaillé de la vue.")

page_break()

# ======================================================================
h1("9. Choix de l'algorithme : pourquoi BFS")
p("Le chapitre 2 a défini graphe, DAG et familles ; le chapitre 8 a décrit "
  "l'algorithme retenu. Ce chapitre caractérise le graphe réel, formule le "
  "problème, puis compare BFS point par point à chaque autre algorithme "
  "classique de graphe.")

h2("9.1 Le graphe manipulé par l'application")
table([
    ["Caractéristique", "RestitutionGraphe", "Conséquence"],
    ["Orienté", "oui (colonne Direction)", "le parcours respecte le sens des arêtes"],
    ["Pondéré", "non", "BFS donne le plus court chemin ; Dijkstra inutile"],
    ["Acyclique (DAG)", "non garanti", "il faut marquer les sommets visités"],
    ["Densité", "creux : 2 à 6 arêtes sortantes par nœud",
     "listes d'adjacence ; BFS peu coûteux par palier"],
    ["Taille", "~100 000 nœuds, centaines de milliers d'arêtes",
     "on ne charge jamais tout le graphe ; parcours borné, par lots en SQL"],
    ["Connexité", "non garantie", "beaucoup de couples sans aucun chemin — "
     "cas fréquent, à traiter efficacement (voir § 11.4)"],
], font=9)

h2("9.2 Le problème à résoudre")
p("Pour un couple (source, cible), l'application répond à deux questions liées :")
numbered([
    "Existence : existe-t-il un chemin orienté de la source vers la cible ?",
    "Plus court chemin : si oui, en fournir un de longueur minimale.",
])
p("La question 1 seule se contente de n'importe quel parcours (BFS ou DFS) : "
  "il suffit d'atteindre la cible. La question 2 impose un parcours par "
  "distance croissante — exactement ce que fait BFS sur un graphe non pondéré. "
  "Résoudre la question 2 résout donc gratuitement la question 1 : on traite "
  "les deux avec un seul BFS.")

h2("9.3 Pourquoi BFS est le meilleur choix ici")
bullet([
    "Graphe non pondéré → BFS est optimal : il rend un plus court chemin en "
    "O(n + m), sans file de priorité ni structure auxiliaire.",
    "Recherche d'existence → BFS s'arrête dès qu'il touche la cible ; sur un "
    "couple relié par un chemin court, il n'explore qu'un petit voisinage.",
    "Réponses « aucun chemin » fréquentes → BFS borné en profondeur (maxDepth) "
    "et en sommets (30 000 par sens) donne une réponse négative en temps "
    "maîtrisé.",
    "Exécution par paliers → un niveau de BFS = une requête SQL par lot ; le "
    "parcours épouse l'accès distant aux données (pas une requête par arête).",
    "Symétrie source/cible → le BFS se dédouble en BFS bidirectionnel "
    "(chapitre 8), qui divise par deux la profondeur explorée : gain "
    "exponentiel sur les chemins longs.",
])
p("En résumé : BFS est le seul algorithme qui donne le plus court chemin d'un "
  "graphe non pondéré au prix d'un simple parcours. Les autres (ci-dessous) "
  "sont soit équivalents mais sans garantie de plus court chemin (DFS), soit "
  "plus généraux mais plus coûteux car ils traitent un problème que nous "
  "n'avons pas — des poids, ou toutes les paires à la fois.")

h2("9.4 Les autres algorithmes, comparés à BFS")

h3("9.4.1 DFS — parcours en profondeur")
p("Principe : explorer aussi loin que possible le long d'une branche avant de "
  "revenir en arrière (pile explicite ou récursion).")
bullet([
    "Résout : existence d'un chemin, détection de cycles, tri topologique, "
    "composantes (fortement) connexes.",
    "Poids : non exploités.",
    "Complexité : O(n + m) — identique à BFS.",
    "Plus court chemin : NON. Le premier chemin que DFS trouve vers la cible "
    "peut être arbitrairement long.",
    "vs BFS : même coût, mais DFS n'ordonne pas les sommets par distance. Pour "
    "« existe-t-il un chemin ? » les deux conviennent ; dès qu'on veut le plus "
    "court, seul BFS répond. DFS s'enfonce volontiers dans une branche "
    "profonde inutile, ce que la borne maxDepth du BFS évite.",
    "Pertinence ici : utilisé pour des besoins annexes (cycles, composantes — "
    "voir § 11.4/11.5), pas pour la recherche de chemin.",
])

h3("9.4.2 BFS bidirectionnel")
p("Principe : deux BFS simultanés, l'un depuis la source (arêtes sortantes), "
  "l'autre depuis la cible (arêtes entrantes), qui se rejoignent au milieu.")
bullet([
    "Résout : plus court chemin en graphe non pondéré, comme BFS.",
    "Complexité : O(n + m) au pire ; en pratique ≈ 2·d^(R/2) sommets visités "
    "au lieu de d^R pour un BFS simple (d = degré moyen, R = longueur du "
    "chemin).",
    "Plus court chemin : OUI (avec alternance correcte des deux fronts).",
    "vs BFS simple : même résultat, coût très inférieur sur les chemins longs, "
    "au prix d'un peu plus de code (deux dictionnaires, condition de "
    "rencontre). C'est la version retenue (chapitre 8).",
])

h3("9.4.3 Dijkstra")
p("Principe : un BFS « à priorité » — on étend toujours le sommet non fixé "
  "dont la distance provisoire (somme des poids) est la plus faible, via une "
  "file de priorité (tas).")
bullet([
    "Résout : plus court chemin dans un graphe pondéré à poids positifs ou nuls.",
    "Poids : requis ≥ 0. Poids négatifs → résultat faux.",
    "Complexité : O((n + m) log n) avec un tas binaire.",
    "Plus court chemin : OUI, en présence de poids.",
    "vs BFS : Dijkstra EST la généralisation de BFS aux poids. Sur un graphe "
    "non pondéré (tous poids = 1), Dijkstra et BFS renvoient le même chemin, "
    "mais Dijkstra paie en plus le coût de la file de priorité (facteur log n) "
    "pour rien. L'employer ici serait correct, mais plus lent et plus "
    "compliqué, sans bénéfice.",
    "Pertinence ici : aucune. Ce serait le bon choix si LINE_VIS_EDG portait "
    "un coût par arête.",
])

h3("9.4.4 Bellman-Ford")
p("Principe : relâcher toutes les arêtes n-1 fois de suite (mise à jour "
  "répétée des distances provisoires).")
bullet([
    "Résout : plus court chemin pondéré, y compris avec poids négatifs ; "
    "détecte les cycles de poids négatif.",
    "Complexité : O(n·m) — nettement plus lent que Dijkstra et que BFS.",
    "Plus court chemin : OUI.",
    "vs BFS : bien plus général (poids négatifs), bien plus cher. Sans poids, "
    "il n'apporte rien que BFS ne fasse déjà, en O(n·m) au lieu de O(n + m).",
    "Pertinence ici : aucune.",
])

h3("9.4.5 A* (A étoile)")
p("Principe : Dijkstra guidé par une heuristique h(v) estimant la distance "
  "restante de v à la cible ; on étend en priorité les sommets minimisant "
  "« distance parcourue + h ».")
bullet([
    "Résout : plus court chemin point-à-point, quand on dispose d'une bonne "
    "heuristique (typiquement une distance géométrique).",
    "Complexité : au mieux quasi linéaire, au pire celle de Dijkstra ; dépend "
    "entièrement de l'heuristique.",
    "Plus court chemin : OUI si l'heuristique est admissible (ne surestime "
    "jamais la distance réelle).",
    "vs BFS : A* « tire » la recherche vers la cible pour l'accélérer, mais "
    "suppose des poids et une heuristique. Ici les nœuds sont des identifiants "
    "abstraits sans géométrie : aucune heuristique naturelle. Le BFS "
    "bidirectionnel procure la même accélération (tirer depuis les deux bouts) "
    "sans rien supposer.",
    "Pertinence ici : aucune (pas d'heuristique disponible).",
])

h3("9.4.6 Floyd-Warshall")
p("Principe : programmation dynamique calculant les plus courts chemins entre "
  "TOUTES les paires de sommets.")
bullet([
    "Résout : la matrice complète des distances (toutes paires).",
    "Complexité : O(n³) en temps, O(n²) en mémoire.",
    "Plus court chemin : OUI, pour toutes les paires simultanément.",
    "vs BFS : répond à une question bien plus large à un coût rédhibitoire. "
    "Ici n peut valoir 100 000 : n³ = 10^15 opérations, impossible. BFS "
    "répond à la seule paire demandée.",
    "Pertinence ici : aucune (mais l'idée de pré-calculer l'atteignabilité, en "
    "plus compact, revient au § 11.6).",
])

h3("9.4.7 Union-Find (composantes connexes)")
p("Principe : structure de données regroupant les sommets par composante au "
  "fur et à mesure de l'ajout des arêtes ; répond en temps quasi constant à "
  "« u et v sont-ils dans la même composante ? ».")
bullet([
    "Résout : la connectivité NON orientée en masse.",
    "Complexité : ~O(m·α(n)) pour construire, ~O(1) par requête.",
    "Plus court chemin : NON — il dit seulement s'il existe un chemin, et "
    "seulement en non orienté.",
    "vs BFS : plus rapide pour répondre en masse à « existe-t-il un chemin ? », "
    "MAIS (1) il faut d'abord parcourir toutes les arêtes pour le construire, "
    "or on ne charge jamais tout le graphe ; (2) il ignore le sens des arêtes, "
    "alors que la question est orientée ; (3) il ne fournit pas le chemin.",
    "Pertinence ici : pas pour la requête en direct, mais c'est exactement "
    "l'outil du pré-calcul des composantes proposé au § 11.4.",
])

h3("9.4.8 Tri topologique (plus court chemin en DAG)")
p("Principe : ordonner les sommets d'un DAG pour que toute arête aille « vers "
  "l'avant », puis relâcher les arêtes dans cet ordre.")
bullet([
    "Résout : plus court (ou plus long) chemin dans un graphe orienté SANS "
    "cycle, en O(n + m), même avec des poids négatifs.",
    "Plus court chemin : OUI, mais seulement si le graphe est acyclique.",
    "vs BFS : plus puissant que BFS sur un DAG pondéré. Mais il exige "
    "l'absence de cycle — non garantie ici (voir § 2.4) — et n'apporte rien de "
    "plus que BFS en l'absence de poids.",
    "Pertinence ici : non pour le graphe brut ; mais le DAG condensé des "
    "composantes fortement connexes (§ 11.5) est, lui, acyclique.",
])

h2("9.5 Tableau comparatif de synthèse")
table([
    ["Algorithme", "Problème résolu", "Poids", "Complexité", "Plus court chemin ?", "Retenu ?"],
    ["BFS", "parcours / plus court chemin non pondéré", "non", "O(n+m)",
     "OUI (non pondéré)", "base"],
    ["BFS bidirectionnel", "idem, accéléré point-à-point", "non",
     "O(n+m) ; ~2·d^(R/2) en pratique", "OUI", "OUI"],
    ["DFS", "existence, cycles, composantes, tri topo", "n/a", "O(n+m)",
     "NON", "non"],
    ["Dijkstra", "plus court chemin, poids ≥ 0", "≥ 0", "O((n+m) log n)",
     "OUI", "non (pas de poids)"],
    ["Bellman-Ford", "plus court chemin, poids négatifs", "quelconques",
     "O(n·m)", "OUI", "non"],
    ["A*", "plus court chemin point-à-point guidé", "≥ 0 + heuristique",
     "≤ Dijkstra", "OUI (h admissible)", "non (pas d'heuristique)"],
    ["Floyd-Warshall", "plus courts chemins toutes paires",
     "quelconques", "O(n³)", "OUI (toutes paires)", "non (n trop grand)"],
    ["Union-Find", "connectivité non orientée en masse", "n/a", "~O(m·α(n))",
     "NON", "pré-calcul (§ 11.4)"],
    ["Tri topologique", "plus court chemin en DAG", "quelconques", "O(n+m)",
     "OUI (si acyclique)", "sur le DAG condensé (§ 11.5)"],
], font=7.5)
p("Lecture : le graphe étant non pondéré, orienté, potentiellement cyclique et "
  "très grand, tous les algorithmes « pondérés » (Dijkstra, Bellman-Ford, A*, "
  "Floyd-Warshall) résolvent un problème plus dur que le nôtre à un coût "
  "supérieur ; DFS ne donne pas de plus court chemin ; le tri topologique "
  "exige l'acyclicité. Le BFS — dans sa variante bidirectionnelle — est le "
  "seul à répondre exactement à la question posée, au coût minimal. Union-Find "
  "et le tri topologique redeviennent utiles en pré-calcul (chapitre 11).")

page_break()

# ======================================================================
h1("10. Galerie des types de graphes — la fonctionnalité")

h2("10.1 Objet")
p("La galerie (route /Graphes) montre une image par famille de graphe et "
  "rappelle, pour chacune, l'algorithme de plus court chemin adapté. Les "
  "familles elles-mêmes, leurs définitions et leurs illustrations sont au "
  "chapitre 2 ; ce chapitre ne décrit que la mécanique de construction des "
  "images côté serveur.")

h2("10.2 Construction des images (SvgGraphRenderer)")
p("Aucune bibliothèque de dessin : le service calcule lui-même les positions "
  "et écrit le SVG.")
bullet([
    "Layout circulaire : chaque composante connexe est placée sur son propre "
    "cercle ; les cercles sont répartis en grille. Utilisé pour la plupart des "
    "types.",
    "Layout par niveaux : pour l'arbre et le DAG — niveau d'un nœud = distance "
    "depuis une racine (nœud sans prédécesseur pour un graphe orienté, premier "
    "nœud pour un arbre) ; les nœuds d'un même niveau sont alignés.",
    "Tracé : un <circle> + un <text> par nœud, une <line> par arête (avec un "
    "marqueur de flèche si le graphe est orienté), une étiquette encadrée au "
    "milieu de l'arête si le graphe est pondéré.",
])
p("Le catalogue est défini dans Models/GraphSamples.cs, le modèle d'un graphe "
  "d'exemple dans Models/GraphSample.cs, le rendu dans "
  "Services/SvgGraphRenderer.cs (code au chapitre 13).")

page_break()

# ======================================================================
h1("11. Pistes d'amélioration de la base de données")
p("Ce chapitre liste des évolutions possibles. Elles ne changent aucune règle "
  "de gestion : le comportement fonctionnel reste identique, seule la "
  "performance (surtout celle de la recherche d'EXISTENCE d'un chemin) "
  "s'améliore. Comme tout l'accès aux données est isolé dans "
  "LineVisEdgRepository, chacune est un changement localisé.")
p("Les § 11.4 (composantes faibles), § 11.5 (condensation SCC) ET § 11.7 "
  "(graphe en mémoire) sont implémentés dans le dossier dotnet-new-scan/ : "
  "même projet MVC que dotnet-mvc/, plus les services GraphScanService, "
  "SccCondensationService, InMemoryGraphService et une page /Scan. Les § 11.2, "
  "11.3, 11.6, 11.8 restent des recommandations.", italic=True)

h2("11.1 Constat sur la structure actuelle")
bullet([
    "Table unique dbo.LINE_VIS_EDG : le sens d'une arête est encodé par une "
    "colonne Direction, donc chaque requête doit filtrer sur Direction et "
    "dériver la source et la cible.",
    "Pas de table de nœuds : vérifier l'existence d'un nœud demande deux "
    "requêtes (colonne Nodes puis NodesLie).",
    "Rien n'est pré-calculé : chaque recherche repart d'un BFS à froid, avec "
    "un aller-retour SQL par palier. Le coût réel est dominé par la latence de "
    "ces requêtes successives, pas par le calcul en mémoire.",
    "Les recherches les plus chères sont les « aucun chemin » : elles vont "
    "jusqu'au plafond des 30 000 nœuds par sens avant de renoncer.",
])

h2("11.2 Indexation (le minimum indispensable)")
p("Les requêtes de voisinage du BFS doivent s'appuyer sur des index composites "
  "couvrants, sinon elles balaient la table.")
code(
    "CREATE NONCLUSTERED INDEX IX_LVE_Nodes_Direction\n"
    "    ON dbo.LINE_VIS_EDG (Nodes, Direction) INCLUDE (NodesLie);\n"
    "CREATE NONCLUSTERED INDEX IX_LVE_NodesLie_Direction\n"
    "    ON dbo.LINE_VIS_EDG (NodesLie, Direction) INCLUDE (Nodes);\n"
    "-- pour DescribePath (transformation d'une arete donnee) :\n"
    "CREATE NONCLUSTERED INDEX IX_LVE_Nodes_NodesLie_Direction\n"
    "    ON dbo.LINE_VIS_EDG (Nodes, NodesLie, Direction) INCLUDE (Transformation);",
    "SQL — index à créer")
p("Gain : chaque FetchEdgesFrom / FetchEdgesInto passe d'un balayage complet à "
  "une recherche d'index (seek). C'est le prérequis de toutes les autres "
  "pistes.")

h2("11.3 Table d'arêtes normalisée (pré-dériver Direction une fois)")
p("Matérialiser une table dbo.EDGE(SourceId, TargetId) où chaque arête est "
  "déjà stockée dans le bon sens. Les requêtes du BFS deviennent triviales "
  "(plus de filtre Direction, plus de OR, un seul index par sens).")
code(
    "CREATE TABLE dbo.EDGE (\n"
    "    SourceId VARCHAR(8000) NOT NULL,\n"
    "    TargetId VARCHAR(8000) NOT NULL\n"
    ");\n"
    "INSERT INTO dbo.EDGE (SourceId, TargetId)\n"
    "SELECT CASE WHEN Direction = 'predecesseur' THEN Nodes    ELSE NodesLie END,\n"
    "       CASE WHEN Direction = 'predecesseur' THEN NodesLie ELSE Nodes    END\n"
    "FROM dbo.LINE_VIS_EDG;\n"
    "CREATE INDEX IX_EDGE_Source ON dbo.EDGE (SourceId) INCLUDE (TargetId);\n"
    "CREATE INDEX IX_EDGE_Target ON dbo.EDGE (TargetId) INCLUDE (SourceId);\n"
    "-- front avant du BFS :  SELECT TargetId FROM dbo.EDGE WHERE SourceId IN (...)\n"
    "-- front arriere du BFS : SELECT SourceId FROM dbo.EDGE WHERE TargetId IN (...)",
    "SQL — table d'arêtes normalisée")
p("À tenir à jour par déclencheur (trigger) sur LINE_VIS_EDG ou par recalcul "
  "périodique. Gain : requêtes plus simples et plus rapides, code du Model "
  "allégé (deux méthodes FetchEdges au lieu de quatre requêtes).")

h2("11.4 Pré-calcul des composantes connexes — le « scan » unique")

h3("11.4.1 L'intuition : des îles")
p("On peut voir le graphe comme un archipel. Les nœuds sont des maisons, les "
  "arêtes sont des ponts. Un groupe de maisons reliées entre elles par des "
  "ponts forme une île — une composante connexe. Deux îles n'ont aucun pont "
  "entre elles.")
p("Constat : si la source est sur une île et la cible sur une autre, il "
  "n'existe AUCUN chemin entre elles — inutile de chercher, il n'y a pas de "
  "pont entre les îles. C'est vrai quel que soit le sens des arêtes (on "
  "raisonne ici en connexité faible).")
p("Exemple (jeu de test de dotnet-new-scan/, script "
  "scripts/seed_disconnected_test.sql) — trois îles :")
table([
    ["Île (ComponentId)", "Nœuds"],
    ["0", "N1, N2, … N100000  (le graphe d'origine)"],
    ["1", "X1, X2, X3, X4, X5"],
    ["2", "Y1, Y2, Y3"],
], font=9.5)
p("Une recherche N1 → X1 traverse deux îles (0 et 1) : réponse « aucun "
  "chemin », immédiate. Une recherche X1 → X5 reste sur l'île 1 : le BFS est "
  "lancé.")

h3("11.4.2 Ce que le scan calcule et stocke")
p("Le scan fait UN seul balayage complet de la table et attribue à chaque "
  "nœud son numéro d'île (ComponentId), stocké dans dbo.NODE_COMPONENT.")
code(
    "CREATE TABLE dbo.NODE_COMPONENT (\n"
    "    NodeId      VARCHAR(450) NOT NULL PRIMARY KEY,  -- <= 900 octets (cle d'index)\n"
    "    ComponentId INT          NOT NULL\n"
    ");\n"
    "CREATE INDEX IX_NODE_COMPONENT_Comp ON dbo.NODE_COMPONENT (ComponentId);",
    "SQL — table des composantes")

h3("11.4.3 Comment il regroupe les nœuds : Union-Find")
p("Au départ, chaque nœud est sa propre île. On lit les arêtes une par une ; "
  "chaque arête u—v dit « u et v sont reliés » → on fusionne leurs deux îles. "
  "La structure Union-Find (§ 9.4.7) fait ces fusions en temps quasi constant "
  "(elle retient, pour chaque nœud, un pointeur vers un « chef d'île »).")
p("Déroulé sur l'île X — arêtes lues dans l'ordre X1—X2, X2—X3, X3—X1, "
  "X3—X4, X4—X5 :")
table([
    ["Arête lue", "Îles après fusion"],
    ["(départ)", "{X1} {X2} {X3} {X4} {X5}"],
    ["X1—X2", "{X1,X2} {X3} {X4} {X5}"],
    ["X2—X3", "{X1,X2,X3} {X4} {X5}"],
    ["X3—X1", "{X1,X2,X3} {X4} {X5}   (déjà ensemble : rien à faire)"],
    ["X3—X4", "{X1,X2,X3,X4} {X5}"],
    ["X4—X5", "{X1,X2,X3,X4,X5}"],
], font=9.5)
p("À la fin, une seule île : elle reçoit un numéro, et l'on écrit X1→1, "
  "X2→1, …, X5→1 dans NODE_COMPONENT.")
p("Union-Find retient, pour chaque nœud, un pointeur parent ; le « chef "
  "d'île » est le nœud dont le parent est lui-même. Deux opérations :",
  bold=True)
code(
    "Find(x)  : remonter les pointeurs parent jusqu'a la racine (le chef).\n"
    "Union(a, b) : ra = Find(a) ; rb = Find(b) ;\n"
    "              si ra != rb, faire pointer l'une des racines vers l'autre.\n"
    "\n"
    "// implementation C# (dotnet-new-scan/Services/GraphScanService.cs)\n"
    "public string Find(string x) {\n"
    "    Ensure(x);                       // 1er contact : x est sa propre racine\n"
    "    while (_parent[x] != x) {\n"
    "        _parent[x] = _parent[_parent[x]];   // compression de chemin\n"
    "        x = _parent[x];\n"
    "    }\n"
    "    return x;\n"
    "}\n"
    "public void Union(string a, string b) {\n"
    "    var ra = Find(a); var rb = Find(b);\n"
    "    if (ra == rb) return;\n"
    "    if (_rank[ra] < _rank[rb]) (ra, rb) = (rb, ra);  // union par rang\n"
    "    _parent[rb] = ra;\n"
    "    if (_rank[ra] == _rank[rb]) _rank[ra]++;\n"
    "}",
    "Union-Find")
p("Deux optimisations rendent chaque opération quasi O(1) amortie : la "
  "compression de chemin (Find aplatit l'arbre au passage) et l'union par "
  "rang (on accroche toujours la petite île sous la grande).")

h3("11.4.4 Étapes 2 et 3 — numéroter, puis comparer")
p("Étape 2 — attribuer un entier à chaque île et remplir NODE_COMPONENT :",
  bold=True)
code(
    "componentOfRoot <- {}   ; id_suivant <- 0\n"
    "pour chaque noeud du graphe :\n"
    "    root <- Find(noeud)\n"
    "    si root absent de componentOfRoot :\n"
    "        componentOfRoot[root] <- id_suivant ; id_suivant <- id_suivant + 1\n"
    "    ecrire (noeud, componentOfRoot[root]) dans NODE_COMPONENT   // via SqlBulkCopy",
    "Pseudo-code — numérotation")
p("Étape 3 — à chaque recherche, avant le BFS :", bold=True)
code(
    "ile_source <- SELECT ComponentId FROM NODE_COMPONENT WHERE NodeId = @source\n"
    "ile_cible  <- SELECT ComponentId FROM NODE_COMPONENT WHERE NodeId = @cible\n"
    "\n"
    "si ile_source est NULL ou ile_cible est NULL :\n"
    "    lancer le BFS            // scan pas fait, ou noeud inconnu : le scan ne casse rien\n"
    "sinon si ile_source != ile_cible :\n"
    "    retourner « aucun chemin »   // O(1), SANS BFS\n"
    "sinon :\n"
    "    lancer le BFS bidirectionnel comme d'habitude",
    "Pseudo-code — utilisation")
p("Sur un graphe non connexe, cela élimine instantanément la majorité des "
  "« aucun chemin » — précisément les recherches les plus coûteuses "
  "aujourd'hui.")

h3("11.4.5 Limite et coût")
p("La connexité FAIBLE ne prouve pas l'existence d'un chemin ORIENTÉ : deux "
  "nœuds d'une même île peuvent ne pas être reliés en respectant le sens des "
  "arêtes (ex. X5 → X1 : même île, mais aucun chemin orienté — le BFS le "
  "confirme). Le scan donne donc un NON certain quand les îles diffèrent, et "
  "un « peut-être » sinon. Pour un OUI/NON orienté exact sans BFS, voir § 11.5.")
table([
    ["Opération", "Coût"],
    ["scan complet (une fois)", "O(n + m) — mesuré ~2,2 s / 100 000 nœuds / 400 000 arêtes"],
    ["comparaison par recherche", "O(1) — 2 lectures sur clé primaire"],
], font=9.5)

h3("11.4.6 Le code (dotnet-new-scan/) — séparation service / repository")
p("Implémentation : dotnet-new-scan/ reprend dotnet-mvc/ et ajoute la page "
  "/Scan (ScanController), la consultation du scan dans HomeController.Index "
  "avant le BFS, et surtout la séparation suivante — le SERVICE ne contient "
  "que l'algorithme, les REQUÊTES SQL sont toutes dans des repositories :")
table([
    ["Élément", "Rôle", "SQL ?"],
    ["Services/GraphScanService.cs", "Union-Find, orchestration des 3 étapes, "
     "verdict de comparaison", "aucune"],
    ["Models/LineVisEdgRepository.cs", "+ StreamAllEdges() : SELECT Nodes, "
     "NodesLie FROM LINE_VIS_EDG (en flux)", "oui"],
    ["Models/NodeComponentRepository.cs", "ReplaceAll(rows) : DROP/CREATE + "
     "SqlBulkCopy + index ; GetComponentIds(a, b) : lecture", "oui"],
], font=9)
p("Le service, débarrassé de SqlConnection / SqlCommand :", bold=True)
code(
    "public ScanStatus Run() {\n"
    "    var uf = new UnionFind();\n"
    "    long edgeCount = 0;\n"
    "\n"
    "    // 1. arêtes fournies par le repository ; le service se contente de fusionner\n"
    "    foreach (var (from, to) in _edges.StreamAllEdges()) {\n"
    "        uf.Union(from, to);\n"
    "        edgeCount++;\n"
    "    }\n"
    "\n"
    "    // 2. un entier par racine -> liste (NodeId, ComponentId)\n"
    "    var componentOfRoot = new Dictionary<string, int>();\n"
    "    var rows = new List<(string NodeId, int ComponentId)>(uf.Nodes.Count);\n"
    "    foreach (var node in uf.Nodes) {\n"
    "        var root = uf.Find(node);\n"
    "        if (!componentOfRoot.TryGetValue(root, out var id)) {\n"
    "            id = componentOfRoot.Count; componentOfRoot[root] = id;\n"
    "        }\n"
    "        rows.Add((node, id));\n"
    "    }\n"
    "\n"
    "    // 3. persistance : entierement deleguee au repository\n"
    "    _components.ReplaceAll(rows);\n"
    "\n"
    "    return new ScanStatus(/* date, n, edgeCount, nb composantes, duree, ... */);\n"
    "}\n"
    "\n"
    "public ComponentVerdict Compare(string source, string target) {\n"
    "    var (cs, ct) = _components.GetComponentIds(source, target);\n"
    "    if (cs is null || ct is null) return ComponentVerdict.ScanUnavailable;\n"
    "    return cs != ct ? ComponentVerdict.DifferentComponents\n"
    "                    : ComponentVerdict.SameComponent;\n"
    "}",
    "C# — GraphScanService (aucune requête SQL)")
p("Les fichiers complets sont dans le dépôt : "
  "dotnet-new-scan/Services/GraphScanService.cs, "
  "dotnet-new-scan/Models/NodeComponentRepository.cs.")

h2("11.5 Condensation en composantes fortement connexes (SCC)")
p("Le scan du § 11.4 ne donne qu'un « non » certain quand les nœuds sont dans "
  "des composantes FAIBLES différentes. La condensation SCC va plus loin : un "
  "verdict d'existence orientée EXACT (oui comme non), sans BFS sur le graphe "
  "d'origine.")

h3("11.5.1 Principe")
numbered([
    "Calculer les composantes fortement connexes (SCC) : sous-ensembles "
    "maximaux de nœuds deux à deux reliés par un chemin orienté dans chaque "
    "sens. Algorithme de Kosaraju — deux parcours en profondeur, O(n + m).",
    "Contracter chaque SCC en un super-nœud. Le graphe des super-nœuds — le "
    "« graphe condensé » — est TOUJOURS un DAG (§ 2.4), et beaucoup plus petit.",
    "Stocker NodeId → SccId (dbo.NODE_SCC) et les arêtes du DAG condensé "
    "SccId → SccId (dbo.SCC_EDGE).",
])
p("Existence d'un chemin orienté u → v :", bold=True)
bullet([
    "SccId(u) = SccId(v) → OUI (même SCC = mutuellement atteignables).",
    "sinon : OUI ⟺ SccId(v) est atteignable depuis SccId(u) par un BFS SUR LE "
    "DAG CONDENSÉ — quelques milliers de super-nœuds, pas 100 000.",
])

h3("11.5.2 Kosaraju (les deux parcours)")
code(
    "// 1. post-ordre : DFS iteratif sur G, empiler les noeuds dans l'ordre\n"
    "//    ou leur exploration se termine.\n"
    "order = PostOrder(G)\n"
    "\n"
    "// 2. DFS iteratif sur G transpose (aretes inversees), en prenant les\n"
    "//    noeuds dans l'ordre inverse de `order`. Chaque arbre de parcours\n"
    "//    est une SCC.\n"
    "sccId = tableau[n] ; scc = 0\n"
    "pour i de order.Count-1 a 0 :\n"
    "    s = order[i]\n"
    "    si s deja visite : continuer\n"
    "    DFS(s sur G_transpose) : chaque noeud atteint recoit sccId = scc\n"
    "    scc = scc + 1\n"
    "\n"
    "// 3. graphe condense : une arete (sccId[u] -> sccId[v]) par arete u->v\n"
    "//    de G dont les extremites sont dans des SCC differentes (dedupliquee).",
    "Pseudo-code — Kosaraju + graphe condensé")
p("Les deux DFS sont ITÉRATIFS (pile explicite) : une version récursive "
  "déborderait la pile d'appel sur 100 000 nœuds.")

h3("11.5.3 Résultat mesuré (dotnet-new-scan/)")
table([
    ["Grandeur", "Graphe d'origine", "Graphe condensé"],
    ["Nœuds", "100 000", "1 979 super-nœuds"],
    ["Arêtes", "~400 000", "2 144"],
    ["Structure", "cyclique", "DAG"],
    ["Plus grande SCC", "—", "98 028 nœuds (SCC géante)"],
], font=9)
p("Un graphe orienté aléatoire dense a typiquement une SCC géante ; "
  "l'atteignabilité sur le DAG condensé (2 000 nœuds) est alors quasi "
  "instantanée.")

h3("11.5.4 Gain concret par rapport au § 11.4")
p("Cas que le scan des composantes faibles ne pouvait pas trancher : X5 → X1 "
  "(dans le jeu de test — {X1,X2,X3} forment un cycle donc une SCC, X4 et X5 "
  "sont des SCC singleton ; le DAG condensé est {X1,X2,X3} → X4 → X5). X5 et "
  "X1 sont dans la même composante FAIBLE (chemin non orienté), mais X5 est un "
  "puits du DAG condensé : SccId(X1) n'est pas atteignable depuis SccId(X5) → "
  "« aucun chemin », EXACT et immédiat, là où le § 11.4 renvoyait vers le BFS.")

h3("11.5.5 Le code (dotnet-new-scan/)")
code(
    "public SccReach Reachable(string source, string target) {\n"
    "    var (s, t) = _scc.GetSccIds(source, target);          // repository\n"
    "    if (s is null || t is null) return SccReach.Unavailable;\n"
    "    if (s == t) return SccReach.Reachable;                // meme SCC\n"
    "\n"
    "    var adj = _condensed ??= _scc.LoadCondensedAdjacency(); // petit DAG, en cache\n"
    "    var seen = new HashSet<int> { s.Value };\n"
    "    var queue = new Queue<int>(); queue.Enqueue(s.Value);\n"
    "    while (queue.Count > 0) {\n"
    "        var u = queue.Dequeue();\n"
    "        if (u == t.Value) return SccReach.Reachable;\n"
    "        if (adj.TryGetValue(u, out var succ))\n"
    "            foreach (var w in succ) if (seen.Add(w)) queue.Enqueue(w);\n"
    "    }\n"
    "    return SccReach.NotReachable;                         // NON exact, sans BFS\n"
    "}",
    "C# — SccCondensationService.Reachable (aucune requête SQL)")
p("Ordre des vérifications dans HomeController : condensation SCC (exacte) "
  "d'abord ; à défaut, composantes faibles (§ 11.4) ; à défaut, BFS. Quand la "
  "SCC répond « atteignable », le BFS est quand même lancé — uniquement pour "
  "récupérer le tracé du chemin à afficher. Fichiers : "
  "dotnet-new-scan/Services/SccCondensationService.cs, "
  "dotnet-new-scan/Models/SccRepository.cs.")

h2("11.6 Index d'atteignabilité (fermeture transitive)")
bullet([
    "Fermeture transitive complète du DAG condensé : table "
    "REACHABLE(FromScc, ToScc) → réponse d'existence en O(1). Taille "
    "potentiellement quadratique en nombre de SCC : viable seulement si le DAG "
    "condensé reste petit.",
    "Variantes compactes : étiquetage par intervalles sur un arbre couvrant du "
    "DAG (min/max d'un parcours en profondeur), ou 2-hop labeling (GRAIL, PLL) "
    "— réponse quasi O(1) avec un index sous-quadratique, au prix d'une "
    "maintenance plus complexe.",
])

h2("11.7 Graphe en mémoire au démarrage — IMPLÉMENTÉ")
p("Le coût réel du BFS n'est pas le calcul mais les allers-retours SQL : une "
  "salve de requêtes par palier, dont le nombre croît avec la taille des "
  "frontières. À grande échelle (2 000 000 de nœuds) c'est le facteur "
  "limitant. La solution : charger tout le graphe orienté en RAM une fois, et "
  "y exécuter le BFS sans plus jamais toucher la base.")

h3("11.7.1 Représentation — CSR (Compressed Sparse Row)")
p("Les nœuds sont indexés par un entier (0..n-1). Le graphe tient dans quatre "
  "tableaux d'entiers :")
code(
    "_fwdOffset[n+1]  _fwdTarget[m]   ->  successeurs de u = _fwdTarget[_fwdOffset[u] .. _fwdOffset[u+1]]\n"
    "_revOffset[n+1]  _revSource[m]   ->  predecesseurs de u = _revSource[_revOffset[u] .. _revOffset[u+1]]",
    "CSR — adjacence directe et inverse")
table([
    ["Volumétrie", "Tableaux CSR", "+ noms des nœuds"],
    ["100 000 nœuds / 400 000 arêtes", "~4 Mo", "~12 Mo au total (mesuré)"],
    ["2 000 000 nœuds / 8 000 000 arêtes", "~80 Mo", "~300–400 Mo (estimation)"],
], font=9)

h3("11.7.2 Chargement et BFS")
bullet([
    "Au démarrage : GraphPreloader (IHostedService) appelle InMemoryGraphService"
    ".Reload() en tâche de fond — le serveur répond tout de suite, les "
    "premières recherches utilisent le BFS SQL jusqu'à ce que le graphe soit "
    "prêt. Mesuré : 740 ms pour 100 000 nœuds.",
    "Reload() lit LineVisEdgRepository.StreamAllDirectedEdges() (le seul SQL, "
    "dans le repository) et construit le CSR. Rejouable via POST /Scan/ReloadGraph.",
    "DirectedGraph.ShortestPath : le même BFS bidirectionnel, sur les tableaux "
    "CSR, avec des Dictionary<int,int> pour forwardPrev / backwardNext. Aucune "
    "requête. Plafond porté à 300 000 nœuds par sens (plus de latence à "
    "craindre, seulement de la RAM transitoire).",
    "HomeController : si le graphe est chargé, le BFS en mémoire ; sinon, repli "
    "sur le BFS SQL. Résultat identique.",
])

h3("11.7.3 Gain mesuré (dotnet-new-scan/)")
table([
    ["", "BFS SQL (dotnet-mvc/)", "BFS en mémoire (§ 11.7)"],
    ["Coût par palier", "~120 requêtes SQL", "parcours de tableaux"],
    ["Latence d'une recherche", "dizaines à centaines de ms", "~8–10 ms (mesuré)"],
    ["Sensible à la taille du graphe", "oui (frontières → plus de requêtes)", "non (borné par maxDepth / plafond)"],
], font=8.5)
p("Le graphe en mémoire est aussi la base commune des autres pré-calculs "
  "(§ 11.4, § 11.5), qui pourraient s'exécuter dessus au lieu de relire la "
  "table.")
p("Fichiers : dotnet-new-scan/Services/DirectedGraph.cs (CSR + BFS), "
  "dotnet-new-scan/Services/InMemoryGraphService.cs (chargement, statut, "
  "GraphPreloader).")

h2("11.8 Autres pistes")
bullet([
    "Index columnstore sur dbo.EDGE pour les balayages analytiques "
    "(statistiques globales du graphe).",
    "Tables « graph » de SQL Server (NODE / EDGE + clause MATCH) : syntaxe "
    "dédiée aux graphes, mais pas de plus court chemin de longueur arbitraire "
    "performant en standard.",
    "Base de données graphe dédiée (Neo4j, etc.) : recherche de chemin native "
    "et optimisée, au prix d'une brique d'infrastructure supplémentaire et "
    "d'une synchronisation depuis SQL Server.",
])

h2("11.9 Synthèse : quelle piste pour quel objectif")
table([
    ["Objectif", "Piste", "Effort", "Effet sur la recherche d'existence"],
    ["Voisinage d'un nœud plus rapide", "index composites (§ 11.2)", "faible",
     "indispensable — seek au lieu de scan"],
    ["Requêtes BFS simplifiées", "table EDGE normalisée (§ 11.3)", "faible",
     "moyen — requêtes et code allégés"],
    ["NON certain quand îles différentes", "composantes faibles pré-calculées "
     "(§ 11.4) — IMPLÉMENTÉ", "moyen", "très élevé sur graphe non connexe"],
    ["OUI/NON orienté EXACT sans BFS", "condensation SCC (§ 11.5) — IMPLÉMENTÉ",
     "élevé", "maximal — BFS sur le petit DAG condensé"],
    ["Réponse O(1) pré-calculée", "fermeture transitive du DAG condensé "
     "(§ 11.6)", "élevé", "maximal (lecture directe)"],
    ["Supprimer la latence SQL du BFS", "graphe en mémoire (§ 11.7) — IMPLÉMENTÉ",
     "moyen", "très élevé — ~10 ms/recherche au lieu de dizaines à centaines"],
    ["Chemin natif optimisé", "base graphe dédiée (§ 11.8)", "élevé",
     "élevé, mais nouvelle infrastructure"],
], font=8.5)

page_break()

# ======================================================================
h1("12. Exigences non fonctionnelles")

h2("12.1 Performance")
bullet([
    "Paramètres typés VARCHAR(8000) (RG-11) : recherches d'index préservées.",
    "Requêtes séparées par index composite, jamais de OR inter-colonnes.",
    "Lots de 1000 paramètres (RG-12).",
    "BFS bidirectionnel : profondeur effective divisée par deux.",
    "Cache applicatif 5 min (RG-09), y compris les résultats « non trouvé ».",
    "Images de graphes : SVG construit à la main, aucune bibliothèque chargée, "
    "sortie de quelques kilo-octets ; peuvent être figées en fichiers "
    "(/Graphes/Build) et servies en statique.",
    "Pistes d'évolution pour aller plus loin : chapitre 11 (indexation, table "
    "d'arêtes normalisée, pré-calcul des composantes, graphe en mémoire).",
])

h2("12.2 Cache applicatif")
table([
    ["Paramètre", "Valeur"],
    ["Implémentation", "IMemoryCache"],
    ["Clé", "path:{source}:{target}:{effectiveMaxDepth}"],
    ["Durée de vie", "5 minutes"],
    ["Taille d'entrée / capacité", "1 / SizeLimit = 10 000"],
], font=9)

h2("12.3 Sécurité")
bullet([
    "Toutes les requêtes SQL sont paramétrées : pas d'injection via source / "
    "target / id.",
    "Accès base en lecture seule (SELECT uniquement).",
    "Le segment de route {id} de /Graphes/Image ne sert qu'à une recherche "
    "exacte dans un catalogue fixe (aucune valeur libre n'atteint le système "
    "de fichiers).",
    "/Graphes/Build écrit uniquement dans wwwroot/img/graphes/, sous des noms "
    "dérivés des slugs du catalogue.",
    "Auth SQL par Trusted_Connection ; TrustServerCertificate = True (dev local).",
])

h2("12.4 Robustesse et limites")
bullet([
    "Garde-fous : maxDepth ≤ 20 ; 30 000 nœuds visités par sens.",
    "Un dépassement se traduit par « aucun chemin », jamais par une erreur.",
    "Les lignes ramenées dans un lot mais hors frontière sont ignorées.",
])

h2("12.5 Portabilité et configuration")
table([
    ["Variable", "Défaut", "Rôle"],
    ["RESTITUTION_DB_SERVER", r"localhost\SQLEXPRESS01", "serveur SQL"],
    ["RESTITUTION_DB_NAME", "RestitutionGraphe", "nom de la base"],
    ["ASPNETCORE_ENVIRONMENT", "Development", "environnement"],
], font=9)

h2("12.6 Observabilité")
bullet([
    "Journalisation ASP.NET Core standard (appsettings.json).",
    "Page /Graphes/Build : retour visuel des fichiers produits.",
])

page_break()

# ======================================================================
h1("13. Code source — backend C#")
p("Repris tel quel du dépôt (dossier dotnet-mvc/).", italic=True)

for title, rel in [
    ("13.1 PathFinder.RazorMvc.csproj", f"{MVC}/PathFinder.RazorMvc.csproj"),
    ("13.2 Program.cs", f"{MVC}/Program.cs"),
    ("13.3 Controllers/HomeController.cs", f"{MVC}/Controllers/HomeController.cs"),
    ("13.4 Controllers/GraphesController.cs", f"{MVC}/Controllers/GraphesController.cs"),
    ("13.5 Models/LineVisEdgRepository.cs", f"{MVC}/Models/LineVisEdgRepository.cs"),
    ("13.6 Models/PathViewModel.cs", f"{MVC}/Models/PathViewModel.cs"),
    ("13.7 Models/GraphSample.cs", f"{MVC}/Models/GraphSample.cs"),
    ("13.8 Models/GraphSamples.cs", f"{MVC}/Models/GraphSamples.cs"),
    ("13.9 Services/SvgGraphRenderer.cs", f"{MVC}/Services/SvgGraphRenderer.cs"),
    ("13.10 Properties/launchSettings.json", f"{MVC}/Properties/launchSettings.json"),
    ("13.11 appsettings.json", f"{MVC}/appsettings.json"),
]:
    h2(title)
    code(whole(rel), rel.split("/", 1)[1])

page_break()

# ======================================================================
h1("14. Code source — vues Razor et CSS")

for title, rel in [
    ("14.1 Views/_ViewImports.cshtml", f"{MVC}/Views/_ViewImports.cshtml"),
    ("14.2 Views/_ViewStart.cshtml", f"{MVC}/Views/_ViewStart.cshtml"),
    ("14.3 Views/Shared/_Layout.cshtml", f"{MVC}/Views/Shared/_Layout.cshtml"),
    ("14.4 Views/Home/Index.cshtml", f"{MVC}/Views/Home/Index.cshtml"),
    ("14.5 Views/Graphes/Index.cshtml", f"{MVC}/Views/Graphes/Index.cshtml"),
    ("14.6 Views/Graphes/Build.cshtml", f"{MVC}/Views/Graphes/Build.cshtml"),
    ("14.7 wwwroot/css/site.css", f"{MVC}/wwwroot/css/site.css"),
]:
    h2(title)
    code(whole(rel), rel.split("/", 1)[1])

page_break()

# ======================================================================
h1("15. Annexes")

h2("15.1 Matrice de traçabilité")
table([
    ["Exigence / cas", "Composant", "Élément de code"],
    ["UC-01, UC-02, UC-03", "Controller", "HomeController.Index"],
    ["UC-01 (algorithme)", "Model", "LineVisEdgRepository.ShortestPath"],
    ["UC-01 (rencontre)", "Model", "LineVisEdgRepository.BuildBidirectionalPath"],
    ["UC-02 (tableau, RG-13)", "Model", "LineVisEdgRepository.DescribePath / EdgeTransformation"],
    ["RG-01", "Model", "RowExists"],
    ["RG-08, RG-09", "Controller", "HomeController.Index (min, cacheKey, GetOrCreate)"],
    ["RG-10", "Model", "constante maxVisitedPerSide"],
    ["RG-11", "Model", "AddVarChar"],
    ["RG-12", "Model", "Chunks + boucles FetchEdges*"],
    ["§ 4.4", "Model", "ToEdge, FetchEdgesFrom, FetchEdgesInto"],
    ["§ 9 (choix de BFS)", "Model", "LineVisEdgRepository.ShortestPath — BFS "
     "bidirectionnel, adapté au graphe orienté non pondéré"],
    ["UC-04, RG-15", "Controller / Model", "GraphesController.Index, GraphSamples.All"],
    ["UC-05", "Controller / Service", "GraphesController.Image, SvgGraphRenderer.Render"],
    ["UC-06, RG-16", "Controller", "GraphesController.Build"],
    ["§ 10.2 (layouts)", "Service", "SvgGraphRenderer.CircularLayout / LayeredLayout / Components"],
    ["RG-14 (pas d'API)", "Program", "AddControllersWithViews + MapControllerRoute (aucun MapControllers d'API)"],
    ["Servir css / images", "Program", "app.UseStaticFiles()"],
    ["§ 11.4 (composantes faibles)", "dotnet-new-scan/", "GraphScanService "
     "+ NodeComponentRepository — IMPLÉMENTÉ"],
    ["§ 11.5 (condensation SCC)", "dotnet-new-scan/", "SccCondensationService "
     "(Kosaraju) + SccRepository — IMPLÉMENTÉ"],
    ["§ 11.7 (graphe en mémoire)", "dotnet-new-scan/", "DirectedGraph (CSR + BFS) "
     "+ InMemoryGraphService + GraphPreloader — IMPLÉMENTÉ"],
    ["Chapitre 11 (§ 11.2–11.3, 11.6, 11.8)", "—", "recommandations, non implémentées"],
], font=8.5)

h2("15.2 Cas de test de recette")
table([
    ["#", "Entrée", "Attendu"],
    ["T1", "GET /", "200, page avec le seul formulaire"],
    ["T2", "GET /?source=N1&target=N1", "« Un chemin existe (0 arête) » + note longueur 0"],
    ["T3", "N1 et N100 reliés", "chaîne de nœuds + tableau (une ligne par arête, "
     "colonne Transformation renseignée)"],
    ["T4", "cible inexistante", "bandeau rouge « Aucun chemin »"],
    ["T5", "GET /?source=N1&target=", "bandeau « Renseigne un nœud source ET un nœud cible »"],
    ["T6", "maxDepth = 999", "traité comme 20 (RG-08)"],
    ["T7", "GET /Graphes", "200, 10 vignettes"],
    ["T8", "GET /Graphes/Image/arbre", "200, image/svg+xml"],
    ["T9", "GET /Graphes/Image/inconnu", "404"],
    ["T10", "GET /Graphes/Build", "200, 10 fichiers listés ; wwwroot/img/graphes/*.svg créés"],
    ["T11", "grep -r \"api\" sur dotnet-mvc/", "aucune route ni contrôleur d'API"],
], font=8.5)

h2("15.3 Variante dotnet-angular-mvc/ (API + Angular)")
p("Même cœur métier (LineVisEdgRepository, BFS bidirectionnel) mais exposé en "
  "API JSON et consommé par un frontend Angular + Cytoscape. Décrite par la "
  "version 1.0 de ce document. Route équivalente à « / » :")
code(whole("dotnet-angular-mvc/Controllers/PathController.cs"),
     "dotnet-angular-mvc/Controllers/PathController.cs")

# ======================================================================
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(f"OK -> {OUT}")

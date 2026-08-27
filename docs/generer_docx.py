# -*- coding: utf-8 -*-
"""Génère docs/Methodes-graphe-BFS-et-autres.docx : catalogue de toutes les
méthodes de parcours de graphe du dépôt restitutiondonnees, avec leur code."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

REPO = Path(r"C:\Users\amami\GitHub\restitutiondonnees")
OUT = REPO / "docs" / "Methodes-graphe-BFS-et-autres.docx"

doc = Document()

# --- styles de base -------------------------------------------------------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

code_style = doc.styles.add_style("CodeBlock", 1)  # 1 = paragraph style
code_style.font.name = "Consolas"
code_style.font.size = Pt(8.5)
code_style.paragraph_format.space_before = Pt(4)
code_style.paragraph_format.space_after = Pt(10)
code_style.paragraph_format.left_indent = Pt(6)


def _shade(paragraph, fill="F2F2F2"):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)
    # petite bordure gauche
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "9CA3AF")
    pBdr.append(left)
    pPr.append(pBdr)


def code(text, lang=""):
    text = text.strip("\n")
    p = doc.add_paragraph(style="CodeBlock")
    if lang:
        run = p.add_run(f"// {lang}\n" if lang in ("C#", "JavaScript / TypeScript") else f"# {lang}\n")
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    for i, line in enumerate(text.split("\n")):
        r = p.add_run(("\n" if (i or lang) else "") + line.replace("\t", "    "))
    _shade(p)
    return p


def h1(t):
    doc.add_heading(t, level=1)


def h2(t):
    doc.add_heading(t, level=2)


def h3(t):
    doc.add_heading(t, level=3)


def para(t, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.bold = bold
    r.italic = italic
    return p


def kv(label, value):
    p = doc.add_paragraph()
    p.add_run(f"{label} : ").bold = True
    p.add_run(value)
    return p


def bullets(items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def slice_file(relpath, start, end):
    """Lignes start..end (1-indexées, inclus) du fichier — le code réel du dépôt."""
    lines = (REPO / relpath).read_text(encoding="utf-8").splitlines()
    return "\n".join(lines[start - 1:end])


# =======================================================================
# PAGE DE TITRE
# =======================================================================
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Méthodes de parcours de graphe")
r.bold = True
r.font.size = Pt(24)

st = doc.add_paragraph()
st.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st.add_run("BFS et autres algorithmes du dépôt restitutiondonnees\n"
               "Catalogue des méthodes et de leur code source")
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

d = doc.add_paragraph()
d.alignment = WD_ALIGN_PARAGRAPH.CENTER
d.add_run("Généré le 2026-08-27 — code repris tel quel des fichiers du dépôt").italic = True

doc.add_paragraph()

# =======================================================================
h1("Introduction")
para(
    "Le dépôt contient plusieurs implémentations du même domaine métier : la "
    "restitution (visualisation + interrogation) d'un graphe orienté de "
    "transformations de données. On y trouve deux familles d'algorithmes de "
    "parcours, selon l'échelle des données :"
)
bullets([
    "Sur la grosse base SQL Server RestitutionGraphe (table dbo.LINE_VIS_EDG, "
    "jusqu'à ~100 000 nœuds) : recherche de plus court chemin et voisinage à N "
    "sauts, en BFS exécuté « par paliers » directement en SQL, sans jamais "
    "charger tout le graphe. Implémenté en Python (pyodbc) et en C# "
    "(Microsoft.Data.SqlClient).",
    "Sur un petit graphe chargé en mémoire depuis un fichier texte "
    "NOEUDS/ARETES (networkx) : détection de cycles par DFS et calcul d'un "
    "layout hiérarchique par tri topologique (façon Sugiyama). Implémenté en "
    "Python (script, MVC, notebooks).",
])
para(
    "Ce document liste chaque méthode, indique le ou les fichiers qui la "
    "contiennent, décrit son principe et sa complexité, puis reproduit son "
    "code actuel."
)

para("Convention de la table LINE_VIS_EDG (rappel, utilisée par toutes les "
     "méthodes BFS sur SQL) :", bold=True)
bullets([
    "Une ligne = (Nodes, Direction, NodesLie). Pas de table de nœuds séparée : "
    "un nœud est une valeur qui apparaît en colonne Nodes ou NodesLie.",
    "Direction = 'predecesseur'  ->  arête  Nodes -> NodesLie",
    "Direction = 'successeur'    ->  arête  NodesLie -> Nodes",
    "Nodes / NodesLie sont des VARCHAR(8000) : les paramètres SQL sont typés "
    "explicitement en VARCHAR pour que SQL Server fasse une recherche d'index "
    "(seek) et non un balayage (scan).",
])

doc.add_page_break()

# =======================================================================
h1("1. Méthodes BFS (parcours en largeur)")

# ---- 1.1 ----
h2("1.1  shortest_path() — BFS unidirectionnel, SQL par paliers (Python)")
kv("Fichier", "src/restitution/db.py")
kv("Exposé par", "GET /api/path  (src/restitution/api.py, FastAPI)")
para("Principe", bold=True)
bullets([
    "BFS classique à un seul front, partant de la source, sens des arêtes "
    "respecté (arêtes sortantes uniquement).",
    "À chaque palier : on prend toute la frontière courante, on récupère en un "
    "lot SQL toutes ses arêtes sortantes (_fetch_edges_from), on marque les "
    "nouveaux nœuds et on reconstitue le chemin via le dictionnaire prev dès "
    "qu'on atteint la cible.",
    "Bornes : max_depth paliers (défaut 12, plafonné à 20 par l'API) et "
    "max_visited = 30 000 nœuds (garde-fou).",
])
kv("Complexité", "O(nœuds visités + arêtes explorées) ; en pratique degré^R pour "
   "un chemin de longueur R — d'où la version bidirectionnelle en C#.")
code(slice_file("src/restitution/db.py", 349, 387), "Python — src/restitution/db.py")

# ---- 1.2 ----
h2("1.2  LineVisEdgRepository.ShortestPath() — BFS bidirectionnel (C#)")
kv("Fichiers", "dotnet-angular-mvc/Models/LineVisEdgRepository.cs\n"
   "dotnet-angular/backend/LineVisEdgRepository.cs  (identique, au namespace et "
   "aux commentaires près)")
kv("Exposé par", "GET /api/path — PathController.Get (MVC) et app.MapGet (Minimal API)")
para("Principe", bold=True)
bullets([
    "Deux fronts BFS simultanés : l'un avance depuis la source par arêtes "
    "sortantes (FetchEdgesFrom), l'autre depuis la cible par arêtes entrantes "
    "(FetchEdgesInto). On alterne un front par palier (step % 2).",
    "forwardPrev[X] = nœud précédent de X depuis la source ; "
    "backwardNext[X] = nœud suivant de X vers la cible.",
    "Dès qu'un nœud est connu des deux côtés, BuildBidirectionalPath recolle "
    "les deux demi-chemins.",
    "Garde-fou maxVisitedPerSide = 30 000 nœuds par sens ; volontairement "
    "synchrone (un thread du pool par requête HTTP).",
])
kv("Complexité", "~ 2 · degré^(R/2) nœuds visités au lieu de degré^R pour un BFS "
   "à sens unique — gain déterminant sur un graphe dense.")
code(slice_file("dotnet-angular-mvc/Models/LineVisEdgRepository.cs", 172, 271), "C#")
para("Reconstruction du chemin (helper) :", bold=True)
code(slice_file("dotnet-angular-mvc/Models/LineVisEdgRepository.cs", 251, 271), "C#")

# ---- 1.3 ----
h2("1.3  neighborhood() — BFS non orienté par paliers (voisinage à N sauts)")
kv("Fichier", "src/restitution/db.py")
kv("Exposé par", "GET /api/neighborhood")
para("Principe", bold=True)
bullets([
    "BFS à un front, mais NON orienté : à chaque palier on récupère toutes les "
    "lignes où la frontière apparaît en Nodes OU en NodesLie "
    "(_fetch_edges_touching), et on ajoute les deux extrémités de chaque arête.",
    "Borné à la fois en profondeur (depth, 1 à 4) et en nombre de nœuds "
    "(limit, défaut 500) : le drapeau truncated signale qu'on a coupé.",
    "Retourne un sous-graphe { nodes, edges } prêt à dessiner (arêtes "
    "dédupliquées et restreintes aux nœuds réellement visités).",
])
code(slice_file("src/restitution/db.py", 304, 346), "Python — src/restitution/db.py")

doc.add_page_break()

# =======================================================================
h1("2. Fonctions support du BFS (accès aux arêtes)")
para("Ces fonctions ne sont pas des parcours mais sont appelées à chaque palier "
     "des BFS ci-dessus. Le motif commun : deux requêtes SQL séparées (une par "
     "colonne / index) plutôt qu'un OR entre deux colonnes, pour garder des "
     "recherches d'index.")

h2("2.1  _row_exists() / RowExists() — le nœud existe-t-il ?")
code(slice_file("src/restitution/db.py", 186, 195), "Python — src/restitution/db.py")
code(slice_file("dotnet-angular-mvc/Models/LineVisEdgRepository.cs", 78, 90), "C#")

h2("2.2  _fetch_edges_from() / FetchEdgesFrom() — arêtes sortantes de la frontière")
code(slice_file("src/restitution/db.py", 223, 248), "Python — src/restitution/db.py")
code(slice_file("dotnet-angular-mvc/Models/LineVisEdgRepository.cs", 92, 132), "C#")

h2("2.3  FetchEdgesInto() — arêtes entrantes de la frontière (front arrière du BFS bidi)")
code(slice_file("dotnet-angular-mvc/Models/LineVisEdgRepository.cs", 134, 170), "C#")

h2("2.4  _fetch_edges_touching() — arêtes touchant la frontière (BFS non orienté)")
code(slice_file("src/restitution/db.py", 198, 220), "Python — src/restitution/db.py")

h2("2.5  Petits utilitaires : _to_edge / ToEdge, _chunks / Chunks, _ph / _phs")
para("Dérivation (source, cible) d'une ligne brute selon sa Direction :", bold=True)
code(slice_file("src/restitution/db.py", 68, 72), "Python")
code(slice_file("dotnet-angular-mvc/Models/LineVisEdgRepository.cs", 68, 70), "C#")
para("Découpage en lots de paramètres (SQL Server plafonne à ~2100 params) :", bold=True)
code(slice_file("src/restitution/db.py", 63, 65), "Python")
code(slice_file("dotnet-angular-mvc/Models/LineVisEdgRepository.cs", 62, 66), "C#")
para("Placeholders SQL castés en VARCHAR(8000) (Python) / AddVarChar (C#) :", bold=True)
code(slice_file("src/restitution/db.py", 53, 60), "Python")
code(slice_file("dotnet-angular-mvc/Models/LineVisEdgRepository.cs", 55, 60), "C#")

doc.add_page_break()

# =======================================================================
h1("3. Autres méthodes / algorithmes")

h2("3.1  find_back_edges() — DFS de détection de cycles")
kv("Fichiers", "python-script/restitution_graphe.py\n"
   "python-notebook/restitution_graphe.ipynb, python-notebook-pyvis/restitution_graphe.ipynb\n"
   "python-mvc/models/graph_model.py  (variante en cached_property, ci-dessous)")
para("Principe", bold=True)
bullets([
    "DFS récursif classique. Deux ensembles : visited (vus une fois pour "
    "toutes) et on_stack (nœuds actuellement sur la pile d'appel).",
    "Une arête (node -> succ) où succ est déjà on_stack referme un cycle : "
    "c'est une « back edge ».",
    "Sert à has_cycle et à exclure ces arêtes du calcul du layout hiérarchique.",
])
code(slice_file("python-script/restitution_graphe.py", 77, 97), "Python — python-script/restitution_graphe.py")
para("Variante python-mvc (mise en cache, réutilisée par has_cycle et "
     "layered_positions) :", bold=True)
code(slice_file("python-mvc/models/graph_model.py", 77, 104), "Python — python-mvc/models/graph_model.py")

h2("3.2  layered_positions() — tri topologique + affectation de niveaux (Sugiyama)")
kv("Fichiers", "python-script/restitution_graphe.py\n"
   "python-mvc/models/graph_model.py\n"
   "python-notebook/restitution_graphe.ipynb")
para("Principe", bold=True)
bullets([
    "On retire les back edges pour obtenir un DAG, puis nx.topological_sort le "
    "parcourt dans l'ordre.",
    "Niveau d'un nœud = 0 s'il n'a pas de prédécesseur, sinon 1 + max(niveau "
    "des prédécesseurs).",
    "Les nœuds d'un même niveau sont répartis horizontalement (triés) ; y = "
    "-niveau (niveau 0 en haut).",
])
code(slice_file("python-script/restitution_graphe.py", 100, 125), "Python — python-script/restitution_graphe.py")
para("Variante python-mvc (échelles paramétrables, DAG construit à partir du "
     "cache _back_edges) :", bold=True)
code(slice_file("python-mvc/models/graph_model.py", 106, 130), "Python — python-mvc/models/graph_model.py")

h2("3.3  _looks_layerable() — heuristique de densité (choix du layout)")
kv("Fichier", "python-notebook-pyvis/restitution_graphe.ipynb")
para("Au-delà d'environ 1,5 arête par nœud, le graphe est en pratique une "
     "grosse composante fortement connexe : le layout hiérarchique n'a plus de "
     "niveaux clairs et tasse tout. Dans ce cas on bascule sur un layout "
     "physique (forces, solveur barnesHut de vis.js).")
code('''
def _looks_layerable(graph: nx.DiGraph, threshold: float = 1.5) -> bool:
    """Heuristique : au-dela d'une certaine densite d'aretes par noeud, un graphe est
    en pratique une grosse composante fortement connexe (beaucoup de cycles qui se
    recoupent). Un layout hierarchique strict n'a alors plus de niveaux clairs a
    calculer et tasse tous les noeuds sur une poignee de rangees [...]. Un
    layout physique (a base de forces, qui repartit les noeuds en 2D) reste lisible
    dans ce cas."""
    if graph.number_of_nodes() == 0:
        return True
    return graph.number_of_edges() / graph.number_of_nodes() < threshold
''', "Python — python-notebook-pyvis/restitution_graphe.ipynb")

h2("3.4  Construction du graphe : parse_graph() / GraphModel.from_text()")
kv("Fichiers", "python-script/restitution_graphe.py, les deux notebooks (parse_graph)\n"
   "src/restitution/models/graph.py, python-mvc/models/graph_model.py (GraphModel.from_text)")
para("Parse le format texte NOEUDS / ARETES (en-têtes insensibles à la casse et "
     "aux accents, lignes vides et # ignorées) et construit un networkx.DiGraph.")
code(slice_file("python-script/restitution_graphe.py", 43, 74), "Python — python-script/restitution_graphe.py")
para("Version « modèle » (méthode de classe, + from_file) :", bold=True)
code(slice_file("src/restitution/models/graph.py", 33, 70), "Python — src/restitution/models/graph.py")

h2("3.5  build_render_data() — mise en forme pour Cytoscape (View MVC)")
kv("Fichier", "python-mvc/views/graph_view.py")
para("Transforme le GraphModel (positions calculées par layered_positions) en "
     "structure JSON consommée par Cytoscape.js en layout « preset ».")
code(slice_file("python-mvc/views/graph_view.py", 14, 33), "Python — python-mvc/views/graph_view.py")

h2("3.6  Requêtes SQL agrégées (interrogation du graphe, hors parcours)")
kv("Fichier", "src/restitution/db.py")

h3("node_detail() — successeurs et prédécesseurs directs (1 saut)")
code(slice_file("src/restitution/db.py", 251, 301), "Python — src/restitution/db.py")

h3("global_stats() — nombre de nœuds/arêtes, densité, degré moyen")
code(slice_file("src/restitution/db.py", 152, 183), "Python — src/restitution/db.py")

h3("search_nodes() — recherche par sous-chaîne d'identifiant")
code(slice_file("src/restitution/db.py", 75, 88), "Python — src/restitution/db.py")

h3("list_nodes() — page de nœuds distincts + transformations liées")
code(slice_file("src/restitution/db.py", 91, 149), "Python — src/restitution/db.py")

h2("3.7  generate_edges() — génération d'un graphe orienté aléatoire (jeu de test)")
kv("Fichier", "scripts/seed_sqlserver.py")
para("Crée ~100 000 nœuds, 2 à 6 arêtes sortantes par nœud (cibles tirées au "
     "hasard, dédupliquées), pour peupler dbo.LINE_VIS_EDG.")
code(slice_file("scripts/seed_sqlserver.py", 38, 48), "Python — scripts/seed_sqlserver.py")

doc.add_page_break()

# =======================================================================
h1("4. Points d'entrée (API et interface)")

h2("4.1  Route /api/path — côté serveur")
para("FastAPI (Python) :", bold=True)
code(slice_file("src/restitution/api.py", 57, 64), "Python — src/restitution/api.py")
para("Minimal API (C#, dotnet-angular/backend) — avec cache mémoire par "
     "(source, cible, maxDepth) :", bold=True)
code(slice_file("dotnet-angular/backend/Program.cs", 38, 57), "C#")
para("Controller MVC (C#, dotnet-angular-mvc) :", bold=True)
code(slice_file("dotnet-angular-mvc/Controllers/PathController.cs", 27, 46), "C#")

h2("4.2  Recherche de chemin — côté client")
para("web/ (JavaScript pur, appelle l'API FastAPI) :", bold=True)
code(slice_file("web/path.js", 83, 121), "JavaScript / TypeScript")
para("Angular (dotnet-angular / dotnet-angular-mvc) — service HTTP :", bold=True)
code(slice_file("dotnet-angular-mvc/ClientApp/src/app/path-api.service.ts", 19, 37),
     "JavaScript / TypeScript")

doc.add_page_break()

# =======================================================================
h1("5. Tableau récapitulatif")

rows = [
    ("Méthode", "Type", "Fichier(s)", "Rôle"),
    ("shortest_path()", "BFS 1 front", "src/restitution/db.py", "Plus court chemin (Python)"),
    ("ShortestPath()", "BFS bidirectionnel", "dotnet-angular*/**/LineVisEdgRepository.cs", "Plus court chemin (C#)"),
    ("neighborhood()", "BFS non orienté", "src/restitution/db.py", "Voisinage à N sauts"),
    ("_fetch_edges_from / FetchEdgesFrom", "accès arêtes", "db.py / *.cs", "Arêtes sortantes d'une frontière"),
    ("FetchEdgesInto", "accès arêtes", "*.cs", "Arêtes entrantes (front arrière)"),
    ("_fetch_edges_touching", "accès arêtes", "db.py", "Arêtes touchant une frontière"),
    ("_row_exists / RowExists", "requête", "db.py / *.cs", "Existence d'un nœud"),
    ("find_back_edges()", "DFS", "python-script/, python-mvc/, notebooks", "Détection de cycles"),
    ("layered_positions()", "tri topologique", "python-script/, python-mvc/, notebook", "Layout hiérarchique (Sugiyama)"),
    ("_looks_layerable()", "heuristique", "notebook pyvis", "Choix hiérarchique vs physique"),
    ("parse_graph / from_text", "parsing", "graph.py, graph_model.py, scripts, notebooks", "Fichier NOEUDS/ARETES -> DiGraph"),
    ("build_render_data()", "vue", "python-mvc/views/graph_view.py", "GraphModel -> Cytoscape JSON"),
    ("node_detail()", "requête", "src/restitution/db.py", "Voisins directs + transformation"),
    ("global_stats()", "requête agrégée", "src/restitution/db.py", "Nœuds, arêtes, densité, degré"),
    ("search_nodes() / list_nodes()", "requête", "src/restitution/db.py", "Recherche / pagination des nœuds"),
    ("generate_edges()", "génération", "scripts/seed_sqlserver.py", "Graphe aléatoire de démonstration"),
]

table = doc.add_table(rows=len(rows), cols=4)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row in enumerate(rows):
    for j, val in enumerate(row):
        cell = table.cell(i, j)
        cell.text = val
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
                if i == 0:
                    r.font.bold = True

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(f"OK -> {OUT}")
